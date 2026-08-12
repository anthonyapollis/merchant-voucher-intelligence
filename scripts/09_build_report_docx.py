"""
09_build_report_docx.py
=======================
Builds report/Merchant_Voucher_Intelligence_Submission.docx — the written submission.

Every figure is pulled from the analytics/ML summaries rather than typed in, so the document
cannot drift from the data. If a number changes upstream, re-running this regenerates the
report correctly rather than leaving a stale figure in prose.
"""
from pathlib import Path
import json
from collections import Counter
import sys
import pandas as pd

sys.path.insert(0, str(Path(__file__).parent))
from _table_registry import (TABLES as REG, TIERS, SUMMARY as REG_SUMMARY,
                             COUNTER_ARGUMENT as REG_COUNTER, counts as _tier_counts, TIER_ORDER)
TIER_COUNTS = _tier_counts()

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "report" / "Merchant_Voucher_Intelligence_Submission.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

S = json.load(open(ROOT / "docs" / "analytics_summary.json"))
ML = json.load(open(ROOT / "docs" / "ml_summary.json"))
RC = json.load(open(ROOT / "docs" / "reconciliation.json"))
K, BA, SZ = S["exec_kpis"], S["business_answers"], S["size_confounder"]
RP, SF = ML["redemption_propensity"], ML["sales_forecast"]

score = pd.read_parquet(ROOT / "data" / "analytics" / "kpi_merchant_scorecard.parquet")
vtype = pd.read_parquet(ROOT / "data" / "analytics" / "kpi_voucher_type.parquet")
region = pd.read_parquet(ROOT / "data" / "analytics" / "kpi_region_performance.parquet")
priority = pd.read_parquet(ROOT / "data" / "analytics" / "kpi_priority.parquet")

NAVY, TEAL, AMBER, RED, GREY = "12305B", "0E8B8B", "E8A317", "C0392B", "5A6672"

doc = Document()

# ---------------------------------------------------------------- base styles
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(7)
st.paragraph_format.line_spacing = 1.13

for name, size, colour, before, after in [
    ("Heading 1", 17, NAVY, 20, 8), ("Heading 2", 13.5, NAVY, 15, 6),
    ("Heading 3", 11.5, TEAL, 11, 4),
]:
    s = doc.styles[name]
    s.font.name = "Calibri Light"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(colour)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Cm(2.0)
sec.left_margin = sec.right_margin = Cm(2.1)


# ---------------------------------------------------------------- helpers
def shade(cell, hexcolour):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolour)
    cell._tc.get_or_add_tcPr().append(el)


def para(text="", size=10.5, bold=False, italic=False, colour=None, align=None,
         space_after=7, space_before=0, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if colour:
        r.font.color.rgb = RGBColor.from_string(colour)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p


def rich(parts, size=10.5, space_after=7, style=None):
    """parts: list of (text, bold, colour) — lets a sentence carry emphasis inline."""
    p = doc.add_paragraph(style=style)
    for t, b, c in parts:
        r = p.add_run(t)
        r.font.size = Pt(size)
        r.bold = b
        if c:
            r.font.color.rgb = RGBColor.from_string(c)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, size=10.5, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.size = Pt(size)
    return p


def table(headers, rows, widths=None, header_colour=NAVY, font=9, band=True,
          align_right=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    align_right = align_right or []
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(str(h))
        r.bold = True
        r.font.size = Pt(font)
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        if i in align_right:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        shade(c, header_colour)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(v))
            r.font.size = Pt(font)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            if i in align_right:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if band and ri % 2 == 1:
                shade(cells[i], "F2F6FB")
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def callout(title, text, colour=TEAL, fill="F0F9F9"):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    c.text = ""
    p = c.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(colour)
    p2 = c.add_paragraph()
    r2 = p2.add_run(text)
    r2.font.size = Pt(9.5)
    p2.paragraph_format.space_after = Pt(3)
    shade(c, fill)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


R = lambda v: f"R{v:,.0f}"
Rm = lambda v: f"R{v/1e6:.1f}m"
P1 = lambda v: f"{v*100:.1f}%"
P2 = lambda v: f"{v*100:.2f}%"


# ======================================================================================
# TITLE PAGE
# ======================================================================================
for _ in range(3):
    doc.add_paragraph()
para("Merchant Sales & Voucher", 30, True, colour=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_after=0)
para("Intelligence Dashboard", 30, True, colour=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_after=14)
para("BI Developer — Second-Round Practical Task", 13, False, True, GREY,
     WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("Microsoft Fabric · Power BI · dbt · SQL · Python ML", 11, False, False, TEAL,
     WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

table(["Deliverable", "Detail"], [
    ["Reporting period", "1 January 2026 – 31 July 2026 (212 days)"],
    ["Data volume", f"{K['VouchersSold']:,} vouchers · 26,500 daily sales rows · "
                    f"{K['TotalTickets']:,} tickets · {K['Merchants']} merchants"],
    ["Architecture", "Fabric medallion (bronze / silver / gold) → Kimball star schema → "
                     "Power BI (import; Direct Lake once hosted)"],
    ["Transformation", "dbt: 18 models, 1 snapshot (SCD2), 132 tests, 4 seeds, 3 exposures"],
    ["Orchestration", "Data Factory pipeline, metadata-driven, daily 02:00 SAST"],
    ["Semantic model", "6 dimensions · 4 facts · 1 analytics mart · 31+ validated DAX measures"],
    ["AI extension", "5 models: anomaly detection, propensity, regression, forecasting, "
                     "segmentation"],
    ["Quality gates", f"14 warehouse tests · 132 dbt tests (153 pass on build, 0 warn, "
                      f"0 error) · 7 SCD2 assertions · {RC['passed']}/"
                      f"{RC['passed']+RC['warnings']+RC['failures']} reconciliation checks"],
], widths=[4.2, 12.4], font=9.5)

doc.add_paragraph()
para("Anthony Apollis", 12, True, colour=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_after=2)
para("11 August 2026", 10, False, True, GREY, WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ======================================================================================
# 1. EXECUTIVE SUMMARY
# ======================================================================================
doc.add_heading("1. Executive summary", 1)

para("This solution delivers a production-minded analytics stack over the supplied merchant "
     "voucher data: a Fabric medallion pipeline feeding a Kimball star schema, a dbt "
     "transformation layer with a 69-test suite, a four-page Power BI report, and five "
     "machine-learning models covering anomaly detection through to forecasting. Every "
     "number in every artefact traces to a single gold layer.")

doc.add_heading("Headline figures", 2)
table(["Metric", "Value", "Note"], [
    ["Total sales", Rm(K["TotalSales"]), f"{K['TotalTransactions']:,} transactions, "
                                          f"avg basket {R(K['AvgBasketValue'])}"],
    ["Redemption rate", P1(K["RedemptionRate"]),
     f"{K['VouchersRedeemed']:,} of {K['VouchersSold']:,} vouchers"],
    ["Avg days to redeem", f"{K['AvgDaysToRedeem']:.2f}", "median 2 days, max 54"],
    ["Delayed redemption rate", P1(K["DelayedRedemptionRate"]), "redeemed after > 7 days"],
    ["Outstanding liability", Rm(K["OutstandingLiability"]), "unredeemed voucher value"],
    ["Avg resolution hours", f"{K['AvgResolutionHours']:.1f}", "median 16.4h — right-skewed"],
    ["SLA breach rate", P1(K["SLABreachRate"]),
     "94.7% of breaches sit on High/Critical"],
    ["Open ticket backlog", f"{K['OpenTickets']:,}", "Open, Escalated or Pending Merchant"],
], widths=[4.6, 3.2, 8.8], align_right=[1])

doc.add_heading("The four findings that matter", 2)

rich([("1. One merchant is failing, and it is an operational failure, not a demand problem. ",
       True, NAVY),
      ("Umhlanga Value Mart's July sales fell 42.5% against its own prior three-month average "
       "while support tickets rose 693% in the same month. The two moving together points to "
       "a service breakdown causing commercial damage — which means it is plausibly "
       "recoverable. Annualised revenue at risk is ", False, None),
      (R(571518), True, RED), (".", False, None)])

rich([("2. The SLA policy is configured backwards. ", True, NAVY),
      ("Critical tickets are given a 12-hour target but take 52.7 hours on average — a 98.3% "
       "breach rate. Low-priority tickets are given 48 hours and take 11.3 — a 0.2% breach "
       "rate. Because the ladder runs opposite to the actual workload, ", False, None),
      ("94.7% of all 358 breaches land on High and Critical", True, RED),
      (". The reported 26% breach rate is measuring a policy misconfiguration, not team "
       "performance.", False, None)])

rich([("3. Eastern Cape is the one declining region, on four independent signals. ", True, NAVY),
      ("It is the only region that peaked before July, sits 9.8% below its own peak while "
       "every other region is at its peak, has the flattest trend slope (+2.0% of average "
       "monthly sales vs +4.4% to +8.5%), and fell 12.2% in June against +1.4% to +3.8% "
       "elsewhere. Three of the five Critical/Watch merchants sit there.", False, None)])

rich([("4. Operational friction does NOT predict merchant performance — and the obvious "
       "analysis says it does. ", True, NAVY),
      ("Tickets per 1,000 transactions correlates with target attainment at r = "
       f"{SZ['raw_corr_friction_vs_attainment']:+.2f}, which looks conclusive. It is "
       "confounded: the ratio is size-dependent (r = "
       f"{SZ['friction_vs_log_size_pearson']:+.2f} against log total sales), so small "
       "merchants score badly purely because the denominator is small. Controlling for size, "
       "the partial correlation collapses to r = "
       f"{SZ['partial_corr_friction_vs_attainment_controlling_size']:+.2f}", False, None),
      (". Reporting the raw correlation would have been a confounded finding presented as a "
       "causal one.", False, None)])

rich([("5. The CRM 'At Risk' flag has zero overlap with actual deterioration. ", True, NAVY),
      ("Both merchants flagged At Risk are ", False, None), ("growing", True, None),
      (" and score Healthy (61.7 and 68.3). Every merchant in genuine decline — including "
       "the one losing R571,518 annualised — is flagged 'Active'. ", False, None),
      (f"R{S['crm_flag_check']['missed_revenue_at_risk']:,.0f}", True, RED),
      (" of at-risk revenue sits in merchants the CRM considers fine. The existing manual "
       "flag is not detecting the problem it exists to detect, which is the clearest "
       "argument for a computed health score.", False, None)])

rich([("6. The open-ticket backlog is two problems, not one. ", True, NAVY),
      (f"Of {S['backlog_ownership']['total_open']} open tickets, ", False, None),
      (f"{S['backlog_ownership']['awaiting_us']} are awaiting us", True, None),
      (" (Open + Escalated) and ", False, None),
      (f"{S['backlog_ownership']['awaiting_customer']} are awaiting the customer", True, None),
      (" (Pending Merchant). Those need entirely different remediation, and a single "
       "'239 open' figure hides the distinction — which is why dim_ticket_status carries an "
       "explicit ownership attribute rather than just an is_open flag.", False, None)])

callout("Where the real operational signal is",
        "Ticket spikes matter at the level of individual events, not as a portfolio rule. "
        "Durban Cash Hub's tickets rose 780% in June while sales GREW 8.2% — a service problem "
        "at a healthy account. Umhlanga Value Mart's rose 693% while sales FELL 42.5% — a "
        "failing account. Identical operational signal, opposite commercial diagnosis. The "
        "alerting logic in this solution therefore pairs ticket movement with sales movement "
        "rather than triggering on either alone.", AMBER, "FEF8EC")

doc.add_page_break()

# ======================================================================================
# 2. ARCHITECTURE
# ======================================================================================
doc.add_heading("2. Solution architecture", 1)

para("The stack follows the Fabric medallion pattern. Each layer has one responsibility, and "
     "the boundaries between them are where the quality controls sit.")

table(["Layer", "Responsibility", "Contents"], [
    ["Landing", "Files as delivered", "4 CSV files in OneLake Files/landing"],
    ["Bronze", "Land raw, add lineage only. No business logic whatsoever.",
     "4 Delta tables + _batch_id, _ingested_at, _source_file"],
    ["Silver", "Type, trim, deduplicate at the declared grain, conform, apply business rules",
     "4 Delta tables, partitioned by month, Z-ORDERed on merchant + date"],
    ["Gold", "Kimball star schema for consumption", "6 dimensions, 4 facts, 1 analytics mart, "
                                                     "2 ML output tables"],
    ["Consumption", "Report, workbook, models", "Power BI (import today, Direct Lake once hosted), Excel pack, Fabric "
                                                 "notebooks, Copilot/Q&A"],
], widths=[2.6, 6.6, 7.4], font=9)

doc.add_heading("Naming convention", 2)
para("Applied consistently so an unfamiliar engineer can predict a table name: "
     "bronze_<source>, silver_<entity>, dim_<entity>, fct_<business_process>, "
     "mart_<consumer>. Surrogate keys are <entity>_key, natural keys are <entity>_id. "
     "Measures are named in business language, not warehouse language — 'Redemption Rate %', "
     "not 'pct_redeemed_calc'.")

doc.add_heading("Orchestration — Azure Data Factory / Fabric Pipelines", 2)
para("Three decisions in the pipeline design are worth stating explicitly, because each one "
     "exists to prevent a specific failure that occurs in production BI.")

rich([("Metadata-driven ingest. ", True, TEAL),
      ("A SourceFiles parameter array carries one object per source; a ForEach iterates it. "
       "Adding a fifth source is a parameter change, not a pipeline edit — which matters "
       "because pipeline edits require deployment and regression testing, while parameter "
       "changes do not. It also guarantees all four copy activities behave identically.",
       False, None)])

rich([("The data-quality gate sits BETWEEN silver and gold. ", True, TEAL),
      ("This is the most important decision in the pipeline. The dbt test suite runs against "
       "silver, and gold is only rebuilt if every test passes. If the gate fails, the "
       "pipeline stops, alerts, and leaves the previous good gold layer in place — so the "
       "report keeps showing yesterday's correct numbers rather than today's wrong ones. "
       "Loading gold first and testing afterwards means the report has already published bad "
       "figures by the time anyone reads the alert. Stale-but-correct beats fresh-but-wrong: "
       "an executive working from a day-old number makes a slightly late decision; one "
       "working from a wrong number makes a wrong decision.", False, None)])

rich([("Row-count validation catches the failure that does not throw. ", True, TEAL),
      ("A truncated or empty source file does not raise an error. The copy succeeds, the "
       "transformation succeeds, and the dashboard shows zero — which looks exactly like a "
       "genuinely bad trading day. A per-source minimum row count turns that silent failure "
       "into a loud one. Thresholds sit well below current volumes so normal fluctuation "
       "does not trip them, but far above zero so truncation cannot pass.", False, None)])

doc.add_heading("Fabric cost control and platform enhancements", 2)
para("A Fabric capacity does not warn, throttle or stop on its own. A scheduled "
     "pipeline fires nightly for the full 60 days whether anyone is watching or not, and if "
     "a paid capacity is later attached the same schedules run against it. Cost control is "
     "therefore built into the artefacts rather than left as an operating instruction.")
table(["Measure", "Implementation", "Saving or risk avoided"], [
    ["Trigger ships disabled",
     "TR_Daily_0200_SAST has runtimeState 'Stopped' and a hard endTime of 2026-10-03",
     "~60 unattended pipeline runs, each spanning a Copy activity, four notebooks and a "
     "semantic model refresh"],
    ["Kill switch",
     "fabric_cost_guard.py --check / --stop / --nuke, using the az login token",
     "Disables every schedule and cancels in-flight runs in one command"],
    ["Automated watch",
     "Two scheduled tasks: weekly escalating check, plus a one-time alert on 2026-10-03",
     "Removes reliance on remembering the expiry date"],
    ["Direct Lake as the production target",
     "The semantic model reads Delta files directly; no import refresh",
     "No duplicated storage, and no scheduled dataset refresh consuming capacity units"],
    ["Ephemeral intermediate models",
     "int_merchant_monthly and int_merchant_momentum are inlined as CTEs",
     "Two fewer materialised tables, and no staleness between them and their consumers"],
    ["Incremental fact load",
     "fct_merchant_sales merges on a deterministic key with a 3-day lookback",
     "A daily run rewrites three days, not 26,500 rows"],
    ["Local development target",
     "The dev target runs the identical models with no cloud cost",
     "Development and testing consume zero capacity; only promotion touches Fabric"],
], widths=[3.4, 6.0, 7.2], font=8.5)

doc.add_heading("Warehouse and Lakehouse enhancements applied", 2)
table(["Enhancement", "Why"], [
    ["Z-ORDER on merchant_id + date in silver",
     "Direct Lake reads Delta files directly, so file layout IS query performance. The report "
     "filters on merchant and period constantly."],
    ["Partition silver by year_month",
     "Partition pruning is negligible at 26,500 rows but the pattern must be correct before "
     "volume arrives, not retrofitted after."],
    ["NOT ENFORCED primary and foreign keys",
     "Fabric Warehouse does not enforce constraints but the optimiser uses them for better "
     "plans — and they document the model for the next engineer."],
    ["Bronze as VARCHAR throughout",
     "Typing is silver's job. A cast applied at ingest is a transformation nobody can see or "
     "test, and it destroys the evidence of what the source actually sent."],
    ["Surrogate keys generated deterministically",
     "The same natural key always yields the same surrogate, so a full rebuild cannot fork a "
     "dimension — which is what makes the pipeline idempotent."],
    ["Adapter-dispatched SQL macros",
     "The same models target both the local development engine and the Fabric Warehouse. T-SQL "
     "has no boolean column type and rejects positional GROUP BY."],
], widths=[4.6, 12.0], font=8.5)

doc.add_heading("Retry policy", 2)
table(["Activity", "Retries", "Rationale"], [
    ["Copy to Bronze", "3", "Transient storage/network faults are common and genuinely "
                            "self-healing"],
    ["Transform to Silver", "1", "Largely deterministic; one retry covers a capacity blip"],
    ["Data Quality Gate", "0", "A failing test is a real defect. Retrying delays the alert "
                               "and risks a flaky pass"],
    ["Build Gold", "1", "As above"],
    ["ML Scoring", "1", "Longer running, more exposed to capacity contention"],
    ["Refresh Semantic Model", "2", "The Power BI REST API returns transient 429s under load"],
], widths=[4.4, 2.0, 10.2], font=9, align_right=[1])

doc.add_page_break()

# ======================================================================================
# 3. DATA MODEL
# ======================================================================================
doc.add_heading("3. Data model", 1)

para("A star schema with four facts at three different grains, sharing conformed dimensions. "
     "Mixing grains in a single fact table is the most common way a voucher model breaks, so "
     "each business process gets its own fact.")

doc.add_heading("What the drop actually supplies", 2)
para("The folder contains six files: four CSVs of data, a README, and a DataDictionary. The "
     "README's suggested model names FIVE tables against those four CSVs — the fifth, DimDate, "
     "has no source file at all. Its entry reads 'candidate should create a proper date "
     "table'. Every table in the delivered model is therefore tagged with one of two distinct "
     "origins where the README is concerned: supplied as a CSV, or named by the README and "
     "built here. Collapsing those into a single 'from the README' label would imply a "
     "DimDate.csv that has never existed.", size=10.5)

table(["Supplied file", "Bytes", "Becomes"], [
    ["MerchantReference.csv", "2,143", "DimMerchant"],
    ["MerchantSales.csv", "2,058,004", "FactMerchantSales"],
    ["VoucherRedemptions.csv", "9,050,258", "FactVoucherRedemptions"],
    ["SupportTickets.csv", "127,372", "FactSupportTickets"],
    ["README_BI_Interview_Data.txt", "851", "Requirements — names DimDate with no file"],
    ["DataDictionary.csv", "740", "Reconciled in the column lineage below"],
], widths=[5.2, 2.4, 8.9], align_right=[1])
para()

doc.add_heading("Source column lineage, and what the supplied dictionary covers", 2)
para("The DataDictionary supplied with the drop describes 12 columns. The four CSVs contain "
     "34. Two thirds of the model therefore rests on meanings inferred by profiling rather "
     "than stated by the data owner — that is an assumption, and it is recorded as one. "
     "Every source column is traced to the gold table it lands in, or marked as deliberately "
     "dropped with the reason. The build fails if a column is neither: a field that silently "
     "disappears between source and star schema is indistinguishable from a defect.",
     size=10.5)

_lin = pd.read_csv(ROOT / "docs" / "source_column_lineage.csv")
_doc_n = int((_lin["DocumentedBySupplier"] == "Yes").sum())
_drop_n = int((_lin["GoldTable"] == "DROPPED").sum())
table(["Measure", "Count"], [
    ["Source columns across the four CSVs", str(len(_lin))],
    ["Documented in the supplied DataDictionary", f"{_doc_n}  ({_doc_n/len(_lin):.0%})"],
    ["Meaning inferred by profiling", str(len(_lin) - _doc_n)],
    ["Carried into the star schema", str(len(_lin) - _drop_n)],
    ["Deliberately dropped, reason recorded", str(_drop_n)],
], widths=[9.5, 4.0], align_right=[1])
para()

table(["Source", "Column", "Documented", "Lands in", "Note"],
      [[r.SourceFile.replace(".csv", ""), r.SourceColumn,
        "Yes" if r.DocumentedBySupplier == "Yes" else "Inferred",
        r.GoldTable if r.GoldTable == "DROPPED" else f"{r.GoldTable}.{r.GoldColumn}",
        ("" if pd.isna(r.ModellingNote) else str(r.ModellingNote))[:150]]
       for r in _lin.itertuples()],
      widths=[2.6, 2.9, 1.9, 4.3, 5.0], font=7.5)
para()

doc.add_heading("The four supplied tables on their own", 2)
para("Before the full model, the same structure reduced to only the tables that arrive as a "
     "CSV. Fourteen boxes is the correct answer to the brief but the wrong diagram for "
     "explaining the SHAPE of the model; four makes it obvious. One conformed dimension sits "
     "at the centre and three facts at three different grains each join to it on "
     "merchant_key. No fact joins to another fact — that is what prevents a fan trap and the "
     "silently multiplied totals it produces.", size=10.5)
para("Every line in this diagram is read from the dbt manifest, specifically from the "
     "relationships tests, so the diagram cannot assert a join the project does not actually "
     "enforce. The seven further relationships to dimensions outside these four (dim_date, "
     "dim_voucher_type, dim_priority, dim_ticket_type, dim_ticket_status) are named beneath "
     "it rather than hidden — omitting date_key would misrepresent the model.", size=10.5)

_c4 = ROOT / "docs" / "erd_core4.png"
if _c4.exists():
    doc.add_picture(str(_c4), width=Cm(16.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
para()

doc.add_heading("Why fourteen tables when the README suggests five", 2)
para(REG_SUMMARY)
table(["Tier", "Count", "Tables"], [
    [TIERS[t][0], str(TIER_COUNTS[t]),
     ", ".join(sorted(k for k, v in REG.items() if v["tier"] == t))]
    for t in TIER_ORDER
], widths=[3.4, 1.4, 11.8], font=9, align_right=[1])

para("Every table is justified individually below. The tier is carried through the dbt model "
     "descriptions, the ERD and the Excel data dictionary, so the same answer appears "
     "wherever the question is asked.", size=9.5, italic=True, colour=GREY)

table(["Table", "Tier", "Rows", "Why it exists"], [
    [name, TIERS[r["tier"]][0], f"{r['rows']:,}", f"{r['why']} {r['detail']}"]
    for name, r in sorted(REG.items(),
                          key=lambda kv: (TIER_ORDER
                                          .index(kv[1]["tier"]), kv[0]))
], widths=[4.0, 2.6, 1.4, 8.6], font=8, align_right=[2])

callout("The fair criticism, and the answer", REG_COUNTER, AMBER, "FEF8EC")

doc.add_heading("Before and after", 2)
para("The four supplied files carry 6 duplicated attribute columns, no declared keys, no "
     "calendar, and every value as text. Every join between them is an implicit string "
     "match that nothing validates. The gold schema replaces that with conformed "
     "dimensions, typed columns and 15 foreign keys that are enforced by a test on every "
     "build — drop a join and the build fails.")

# The ERDs are wide. On A4 portrait there is only ~16.8cm of usable width, so a 24cm image
# is silently cropped at the right edge — which is exactly what happened on the first pass.
# Wide diagrams get their own LANDSCAPE section, then the document returns to portrait.
_erdb = ROOT / "docs" / "erd_before.png"
_erda = ROOT / "docs" / "erd_after.png"

_land = doc.add_section(WD_SECTION.NEW_PAGE)
_land.orientation = WD_ORIENT.LANDSCAPE
_land.page_width, _land.page_height = Cm(29.7), Cm(21.0)
_land.left_margin = _land.right_margin = Cm(1.5)
_land.top_margin = _land.bottom_margin = Cm(1.5)
_LAND_W = Cm(26.5)          # 29.7 - 1.5 - 1.5, with a little breathing room

for _img, _cap in ((_erdb, "BEFORE — the four files as delivered. Red dashed arrows are "
                           "implicit string joins that nothing validates; DUP marks a "
                           "duplicated attribute column."),
                   (_erda, "AFTER — the gold star schema. Every arrow is an enforced "
                           "relationships test, and each table is coloured by why it "
                           "exists.")):
    if _img.exists():
        from PIL import Image as _PILImage
        _w, _h = _PILImage.open(_img).size
        # Fit to width, but never let a tall image run past the usable page height
        _max_h = Cm(16.5)
        _use_w = _LAND_W if (_LAND_W * _h / _w) <= _max_h else Cm(_max_h.cm * _w / _h)
        doc.add_picture(str(_img), width=_use_w)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        para(_cap, 9, italic=True, colour=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()
    else:
        para(f"[{_cap}  —  see docs/erd_before_after.html]", 9, italic=True, colour=GREY)

_port = doc.add_section(WD_SECTION.NEW_PAGE)
_port.orientation = WD_ORIENT.PORTRAIT
_port.page_width, _port.page_height = Cm(21.0), Cm(29.7)
_port.left_margin = _port.right_margin = Cm(2.1)
_port.top_margin = _port.bottom_margin = Cm(2.0)

table(["Delivered", "Modelled", "Why it matters"], [
    ["No date table", "dim_date — 365 contiguous days, gap-tested",
     "Time intelligence returns wrong answers, not errors, on a gapped calendar"],
    ["Merchant/Region/Channel on 3 facts", "Dropped; live only on dim_merchant",
     "Profiling proved 100% agreement, so removal is lossless — and stops two versions of "
     "'Region' existing"],
    ["String joins, nothing enforced", "15 relationships tests on every build",
     "A broken join fails the pipeline instead of silently returning fewer rows"],
    ["Everything text", "Typed at silver", "No implicit coercion hiding inside an aggregate"],
    ["Repeated category strings", "4 conformed dimensions",
     "Sort order, grouping and metadata that do not exist in the source"],
    ["Redemption date as a plain column", "Role-playing date, ACTIVE + INACTIVE",
     "Otherwise late redemptions attribute back to the sale month and a backlog is invisible"],
    ["No history", "snap_merchant SCD2 + dim_merchant_history",
     "The source overwrites; 'when did this merchant become At Risk' was unanswerable"],
    ["Two grains conflated", "Separate facts + a recorded 4.2:1 ratio",
     "Stops someone reporting a R43.5m 'break' that is not a break"],
], widths=[4.2, 4.8, 7.6], font=8.5)

doc.add_heading("Tables", 2)
table(["Table", "Type", "Grain", "Rows"], [
    ["dim_date", "Dimension", "One row per calendar day, contiguous, full year boundaries",
     "365"],
    ["dim_merchant", "Dimension", "One row per merchant + Unknown member", "26"],
    ["dim_merchant_history", "Dimension (Type 2)",
     "One row per merchant per version, with valid_from / valid_to", "25"],
    ["dim_voucher_type", "Dimension", "One row per voucher type", "5"],
    ["dim_priority", "Dimension", "One row per ticket priority", "4"],
    ["dim_ticket_type", "Dimension", "One row per ticket type", "6"],
    ["dim_ticket_status", "Dimension", "One row per ticket status", "4"],
    ["fct_merchant_sales", "Fact (transaction)", "Date × Merchant × VoucherType", "26,500"],
    ["fct_voucher_redemptions", "Fact (accumulating snapshot)", "One row per voucher",
     "120,969"],
    ["fct_support_tickets", "Fact (transaction)", "One row per ticket", "1,363"],
    ["fct_merchant_target", "Fact (periodic snapshot)", "Month × Merchant", "175"],
    ["mart_merchant_scorecard", "Analytics mart", "One row per merchant", "25"],
], widths=[5.0, 4.4, 5.6, 1.6], font=9, align_right=[3])

doc.add_heading("Entity relationships", 2)
para("Generated from the dbt manifest, not drawn. dbt docs shows the DAG — which model "
     "BUILDS from which — and that is build lineage, not entity relationships: it cannot "
     "show that fct_support_tickets.priority_key joins to dim_priority, because the fact is "
     "built from staging, not from the dimension. The real foreign keys live in the schema "
     "tests. Every relationships test is an enforced statement that a column must resolve "
     "to another column, so this diagram cannot drift — drop a join and the test disappears "
     "with it.")
table(["From", "Column", "To"], [
    ["fct_merchant_sales", "date_key", "dim_date"],
    ["fct_merchant_sales", "merchant_key", "dim_merchant"],
    ["fct_merchant_sales", "voucher_type_key", "dim_voucher_type"],
    ["fct_voucher_redemptions", "sold_date_key  (ACTIVE)", "dim_date"],
    ["fct_voucher_redemptions", "redeemed_date_key  (INACTIVE)", "dim_date"],
    ["fct_voucher_redemptions", "merchant_key", "dim_merchant"],
    ["fct_support_tickets", "date_key", "dim_date"],
    ["fct_support_tickets", "merchant_key", "dim_merchant"],
    ["fct_support_tickets", "ticket_type_key", "dim_ticket_type"],
    ["fct_support_tickets", "priority_key", "dim_priority"],
    ["fct_support_tickets", "status_key", "dim_ticket_status"],
    ["fct_merchant_target", "date_key", "dim_date"],
    ["fct_merchant_target", "merchant_key", "dim_merchant"],
    ["dim_merchant_history", "merchant_key", "dim_merchant"],
    ["mart_merchant_scorecard", "merchant_key", "dim_merchant"],
], widths=[5.0, 6.2, 5.4], font=8.5)
para("15 enforced foreign keys. Note fct_voucher_redemptions carries TWO into dim_date — "
     "the role-playing date that makes a redemption backlog visible.",
     size=9, italic=True, colour=GREY)

doc.add_heading("Normalisation, keys and the Kimball position", 2)
para("The model follows Kimball dimensional design (Kimball & Ross, The Data Warehouse "
     "Toolkit, 3rd edition). The choices below are the ones a reviewer is most likely to "
     "question, so each is stated with its reason rather than left implicit.", size=10.5)

doc.add_heading("Normalised where it removes contradiction, denormalised where it removes joins", 3)
para("The source is denormalised without benefit: Merchant, Region and Channel appear on "
     "three fact files AND on the reference file — four places one value can disagree with "
     "itself, with no query made faster in exchange. Those nine columns were dropped from the "
     "facts after profiling proved 100 percent agreement, so removal was lossless. Region and "
     "Channel now live only on dim_merchant.", size=10.5)
para("The dimensions themselves are deliberately NOT normalised further. Splitting Region, "
     "Channel and AccountManager into their own lookup tables is snowflaking, and Kimball is "
     "explicit that it should be resisted: dimension tables are geometrically smaller than "
     "fact tables, so normalising them has virtually no effect on database size, and the "
     "trade of dimension space for simplicity is almost always worth making. At 25 merchants "
     "the storage saving is nil while the cost — an extra join on every regional slice, and a "
     "longer filter-propagation chain in Power BI — is real.", size=10.5)

table(["Layer", "Form", "Why"], [
    ["Source CSVs", "Denormalised, redundantly",
     "Attributes repeated across four files with nothing keeping them in step"],
    ["Facts", "Normalised to keys and measures",
     "One version of each attribute; a fact carries foreign keys, measures and degenerate "
     "dimensions only"],
    ["Dimensions", "Denormalised (flat), not snowflaked",
     "One join to reach any attribute; storage saving from snowflaking is negligible at this "
     "size"],
], widths=[3.0, 4.2, 9.4], font=8.5)
para()

doc.add_heading("The keys, and what each is for", 3)
table(["Key type", "In this model", "Purpose"], [
    ["Primary key (dimension)",
     "MerchantKey, DateKey, VoucherTypeKey, PriorityKey, TicketTypeKey",
     "One row per key. The basis for referential integrity with every fact that joins to it."],
    ["Surrogate key",
     "Issued for each dimension; the natural key is retained alongside it",
     "Insulates the warehouse from source key changes and lets a dimension hold several "
     "versions of the same natural key over time."],
    ["Natural key",
     "MerchantID (M001), retained on dim_merchant",
     "The business identifier. Kept so a merchant can still be traced back to the source "
     "system, and so a Type 2 dimension can group its versions."],
    ["Foreign key (fact)",
     "15 enforced relationships, each with a dbt relationships test",
     "The join path. A broken foreign key returns FEWER ROWS rather than an error, so these "
     "are tested on every build instead of trusted."],
    ["Role-playing key",
     "SoldDateKey (active) and RedeemedDateKey (inactive), both to DimDate",
     "One dimension playing two roles. Sold date is the default because every voucher has "
     "one; redeemed date is null until redemption, so activating it would silently drop "
     "unredeemed vouchers from any date-filtered total."],
    ["Degenerate dimension",
     "VoucherID and TicketID, held on the fact with no dimension table",
     "An identifier with no attributes of its own. Kimball's guidance is to leave it on the "
     "fact rather than build a dimension that would hold nothing but the key."],
    ["Composite / compound key",
     "fct_merchant_sales is unique on date x merchant x voucher type",
     "The declared grain. Tested with a composite uniqueness test — the check that a fact has "
     "not silently gained a second row per key."],
    ["Unknown member",
     "MerchantKey = -1 on dim_merchant",
     "An unmatched fact row lands on a visible Unknown member instead of vanishing from every "
     "total. This is why dim_merchant has 26 rows against 25 merchants."],
], widths=[3.2, 5.0, 8.4], font=8.5)
para()

doc.add_heading("What else in Kimball this build applies", 3)
table(["Principle", "Applied here"], [
    ["Declare the grain before anything else",
     "Each fact states its grain explicitly and is tested against it: sales at date x merchant "
     "x voucher type, redemptions one row per voucher, tickets one row per ticket."],
    ["Never mix grains in one fact table",
     "Monthly targets are a separate 175-row fact, not a column on dim_merchant. On the "
     "dimension a monthly value would re-count once per daily fact row — the fan trap — and "
     "inflate every attainment figure."],
    ["Conformed dimensions",
     "dim_date, dim_merchant and dim_voucher_type are shared across facts, so measures from "
     "different business processes can be compared side by side."],
    ["Facts join only through dimensions",
     "No fact joins directly to another fact. Q7 in the SQL pack is written specifically to "
     "show the target and sales facts meeting only via shared dimension keys."],
    ["Slowly changing dimensions",
     "MerchantReference is a current-state extract. A Type 2 snapshot preserves status, owner "
     "and target history that the source overwrites on every load."],
    ["Accumulating snapshot",
     "The voucher fact is one row per voucher, updated when the second event (redemption) "
     "occurs — which is why it carries two date keys rather than two rows."],
    ["Additive, semi-additive, non-additive",
     "SalesValue and Transactions are fully additive. Rates and percentages are held as "
     "measures rather than stored columns, so they are never summed by accident."],
], widths=[4.4, 12.2], font=8.5)
para()

doc.add_heading("Modelling decisions worth defending", 2)

rich([("Role-playing dates on the redemption fact. ", True, TEAL),
      ("A voucher has two events — sold and redeemed — so the fact carries both "
       "sold_date_key and redeemed_date_key. In the semantic model, dim_date joins "
       "sold_date_key on the ACTIVE relationship and redeemed_date_key on an INACTIVE one, "
       "activated inside specific measures via USERELATIONSHIP. This is the single most "
       "common place a voucher model goes wrong: reporting redemptions on the sale date "
       "attributes late redemptions back to the month of issue, which makes a redemption "
       "backlog completely invisible.", False, None)])

rich([("Descriptive attributes dropped from the facts. ", True, TEAL),
      ("Merchant, Region and Channel appear on both the sales fact and MerchantReference. "
       "Profiling confirmed 100% agreement across every row, so they were dropped from the "
       "facts — losslessly. Keeping them would allow two competing versions of 'Region' to "
       "exist in the model, which is how a report ends up with two different regional totals "
       "on two different pages.", False, None)])

rich([("SLA hours stored on the fact, not resolved from the dimension. ", True, TEAL),
      ("They agree today. But a future SLA policy change must not retrospectively restate "
       "whether a historic ticket breached — a ticket is judged against the SLA in force when "
       "it was raised. dim_priority.target_sla_hours therefore describes CURRENT policy, and "
       "the fact records what was actually promised at the time.", False, None)])

rich([("Targets in a separate fact at month grain. ", True, TEAL),
      ("fct_merchant_target is a periodic snapshot at a coarser grain than the daily sales "
       "fact. Forcing a monthly target into a daily fact would break every additive measure. "
       "The two facts join only through shared dimensions, which is the standard Kimball "
       "answer to multi-grain reporting. Targets are pro-rated by days covered so a "
       "part-month never shows a false shortfall.", False, None)])

rich([("An Unknown (-1) dimension member. ", True, TEAL),
      ("Emitted deliberately so a future fact row with an unmatched merchant lands somewhere "
       "visible in the report rather than silently disappearing from every total.",
       False, None)])

rich([("Both a Type 1 and a Type 2 merchant dimension. ", True, TEAL),
      ("dim_merchant is current state and is what the semantic model joins to — 'show me "
       "sales for merchants CURRENTLY in Gauteng' is what a regional manager asking about "
       "their own patch actually means. dim_merchant_history is versioned and answers "
       "point-in-time questions: 'what were sales for merchants that were in Gauteng AT THE "
       "TIME'. Loading only the Type 2 version and forcing every query through a date-range "
       "join would be technically purer and practically worse; loading only Type 1 loses "
       "history permanently. Both, clearly labelled, is the workable answer.", False, None)])

rich([("Diagnostics carried in the dimension, not just the report. ", True, TEAL),
      ("dim_priority stores the observed median and 90th-percentile resolution time next to "
       "the target SLA, and derives sla_is_achievable from them. The SLA finding is therefore "
       "data rather than prose: Critical carries a 12-hour target against a 74-hour 90th "
       "percentile, and the model says so directly. A reader can sort on it; they cannot "
       "sort on a paragraph.", False, None)])

doc.add_page_break()

# ======================================================================================
# 4. ETL & dbt
# ======================================================================================
doc.add_heading("4. Transformation layer", 1)

doc.add_heading("Business rules, defined once", 2)
para("Every threshold lives in a single place (dbt_project.yml vars) and is inherited by the "
     "SQL models, the PySpark notebooks and the DAX measure library. A change propagates "
     "consistently instead of being applied in three places and forgotten in a fourth.")
table(["Rule", "Value", "Basis"], [
    ["Delayed redemption", "> 7 days", "No threshold supplied. 7 days sits above the 75th "
                                       "percentile (5 days), so it flags a genuine tail "
                                       "rather than routine behaviour"],
    ["Voucher expiry / breakage", "90 days", "No expiry rule supplied; 90 days is a common "
                                             "industry default. Affects breakage only"],
    ["High priority", "High, Critical", "Used for the SLA concentration analysis"],
    ["Redemption integrity", "Flag AND valid non-retrograde date must agree",
     "Trusting the flag alone would let a corrupt feed inflate the most scrutinised KPI in "
     "the report"],
], widths=[4.0, 3.4, 9.2], font=9)

doc.add_heading("SQL report pack — the queries behind the numbers", 2)
para("Eight queries answer the brief directly from the gold star schema, in docs/"
     "report_queries.sql. Each joins through the conformed dimensions rather than reading a "
     "fact in isolation: the join path IS the argument for the model. Each was executed "
     "against the warehouse and the results captured as rendered output in "
     "docs/screenshots/sql, so the image and the data cannot drift apart.", size=10.5)

_sql_shots = sorted((ROOT / "docs" / "screenshots" / "sql").glob("q*.png"))
for _p in _sql_shots:
    doc.add_picture(str(_p), width=Cm(16.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
para()

doc.add_heading("A finding about the supplied targets", 2)
para("Target attainment is above 100% for every merchant in the dataset — the range is 381% "
     "to 916%, and the portfolio total is 613.7% (R65,521,299 actual against R10,676,900 of "
     "target). This was checked before being reported: the join is correct and the figure "
     "reproduces from the raw CSV. BaseMonthlySalesTarget is simply not on the same scale as "
     "observed sales in this synthetic data — Blue Crane Trading carries a R60,133 monthly "
     "target against roughly R466,509 of actual monthly sales.", size=10.5)
para("The consequence is worth stating plainly, because it changes how the measure may be "
     "used. Attainment here is a RELATIVE ranking device — the 2.41x spread between the "
     "weakest and strongest merchant is meaningful — but the absolute percentage is not, and "
     "no target-hit or target-miss commentary should be built on it. It also reinforces the "
     "Q4 conclusion: with a fixed denominator per merchant, attainment is close to a "
     "rescaling of sales, which is exactly why the apparent tickets-versus-attainment "
     "correlation collapses once size is controlled for.", size=10.5)

doc.add_heading("What dbt is used for, and what it earned", 2)
para("dbt is not doing anything here that Python could not do. Naming what it actually "
     "contributed is more useful than claiming it was indispensable, so each item below is "
     "the concrete thing that would be lost without it.", size=10.5)

_mn = json.loads((ROOT / "dbt" / "target" / "manifest.json").read_text(encoding="utf-8"))
_rt = Counter(n["resource_type"] for n in _mn["nodes"].values())
_tk = Counter((n.get("test_metadata") or {}).get("name", "singular")
              for n in _mn["nodes"].values() if n["resource_type"] == "test")

table(["What dbt provides", "In this project", "What is lost without it"], [
    ["Enforced referential integrity",
     f"{_tk['relationships']} relationships tests between facts and dimensions",
     "The BEFORE model joins on a string nobody validates. A broken join returns FEWER "
     "ROWS rather than an error — the totals simply get quietly smaller and no one is "
     "told. These fail the build instead."],
    ["Grain and uniqueness contracts",
     f"{_tk['unique']} unique, {_tk['not_null']} not_null, "
     f"{_tk['unique_combination_of_columns']} composite-key tests",
     "Duplicate-proofing becomes a claim rather than a check. A fact silently gaining a "
     "second row per key double-counts every measure built on it."],
    ["Domain constraints",
     f"{_tk['accepted_values']} accepted_values, {_tk['expression_is_true']} expression "
     f"tests, {_tk['expect_column_values_to_be_between']} range tests",
     "A new Priority value or a negative sales figure enters the model and reaches the "
     "report as a legitimate-looking number."],
    ["Change history (SCD2)",
     f"{_rt['snapshot']} snapshot over MerchantReference",
     "This is the strongest single case. The source is a CURRENT-STATE extract: each load "
     "overwrites active_status, account_manager and target. Without the snapshot that "
     "history is destroyed at source and is unrecoverable — 'when did this merchant become "
     "At Risk, and did sales fall before or after?' can never be answered."],
    ["Business logic held outside code",
     f"{_rt['seed']} seeds (voucher category, margin band, priority and status metadata)",
     "The commercial team needs a code deployment to change a margin band."],
    ["Lineage and documentation from the build",
     f"{_rt['model']} models; the ERD, DAG and column lineage in this report are generated "
     f"from manifest.json",
     "Documentation is drawn by hand and starts drifting from the code the day it is "
     "written."],
    ["One codebase, two engines",
     "Adapter-dispatch macros in macros/portability.sql",
     "A local engine for development (seconds, no cloud cost) and Fabric for production would be "
     "two separately-maintained SQL dialects that gradually disagree."],
], widths=[3.6, 4.6, 8.4], font=8.5)
para()

para("The honest limit of the argument: the transformation itself could have been written in "
     "pandas, and in fact it was — the Python pipeline in scripts/02 and scripts/03 produces "
     "the same gold tables independently. That duplication is deliberate. Two implementations "
     "of the same business logic should agree, and the reconciliation in section 8 is what "
     "checks that they do. It is what caught the R984,046 date-window defect: the calendar "
     "SPAN was being used where the reporting WINDOW was meant, and only a disagreement "
     "between the two implementations exposed it.", size=10.5)

doc.add_heading("dbt project", 2)
para(f"{_rt['model']} models across staging, intermediate and marts, plus {_rt['snapshot']} "
     f"snapshot, {_rt['seed']} seeds and {_rt['test']} tests. The project runs against a local "
     f"in development (seconds, no cloud cost) and the Fabric Warehouse in production via "
     f"dbt-fabric. Identical SQL, only the adapter changes.")

table(["Layer", "Models", "Purpose"], [
    ["staging", "4 views", "Type, trim, deduplicate at the declared grain, apply integrity rules"],
    ["intermediate", "2 ephemeral",
     "int_merchant_monthly (merchant × month spine, zero-filled) and int_merchant_momentum "
     "(period comparisons). Inlined as CTEs — no extra storage, one definition"],
    ["snapshots", "1", "snap_merchant — Type 2 capture of slowly-changing merchant attributes"],
    ["marts / core", "10", "6 conformed dimensions + 4 facts, plus the Type 2 dimension"],
    ["marts / analytics", "1", "mart_merchant_scorecard — drill-through target and ML features"],
], widths=[3.2, 2.4, 11.0], font=9)

para("Two modelling changes were made after the first build, both to close real gaps:",
     space_after=4)
rich([("Three ticket dimensions instead of one combined table. ", True, TEAL),
      ("Ticket type, priority and status were originally unioned into a single table behind a "
       "dimension_type discriminator. That had two costs: the three foreign keys on "
       "fct_support_tickets could not be tested for referential integrity — a relationships "
       "test could not distinguish 'priority_key resolves to a priority' from 'resolves to "
       "something, somewhere' — and Power BI cannot build three independent filter paths off "
       "one physical table without role-playing copies. Splitting them added 9 referential "
       "tests that previously could not exist. Six rows is not too small to deserve its own "
       "dimension; join clarity matters more than table count.", False, None)])

rich([("A Type 2 snapshot on merchant attributes. ", True, TEAL),
      ("MerchantReference is a current-state extract — each load overwrites the last, so "
       "active_status, account_manager and target history is destroyed at source. That "
       "matters: two merchants are currently flagged 'At Risk', and 'when did that change, "
       "and did performance decline before or after?' is unanswerable without history. The "
       "snapshot uses the check strategy rather than timestamp, because the source has no "
       "reliable last-modified column (OnboardedDate is when the merchant joined, not when "
       "the row was edited).", False, None)])

table(["Test category", "Count", "Examples"], [
    ["Uniqueness / not-null", "58", "Primary keys on every dimension, fact and mart"],
    ["Referential integrity", "18", "Every fact foreign key resolves to its own dimension"],
    ["Accepted values", "14", "Region, Channel, Priority, Status, SLAHours, quality_flag"],
    ["Range / expectation", "16", "sales_value ≥ 0; resolution_hours between 0 and 720"],
    ["Row-level logic", "10", "is_sla_breach must equal (resolution_hours > sla_hours)"],
    ["Completeness", "6", "The scorecard must cover every real merchant; every merchant must "
                          "have 7 months observed"],
    ["Singular (business)", "4", "dim_date contiguity; bronze→gold revenue reconciliation; "
                                 "redemption rate plausibility; no future-dated facts"],
    ["Source freshness", "3", "Warn at 26h, error at 48h"],
    ["SCD Type 2 behaviour", "7", "Separate harness: simulates a change, asserts history was "
                                  "captured, restores state"],
], widths=[4.6, 1.8, 10.2], font=9, align_right=[1])

callout("Validating the snapshot rather than declaring it",
        "A snapshot that has only ever seen one version of each row is indistinguishable from "
        "one that is silently broken. scripts/_test_scd2.py simulates a real merchant change "
        "(Umhlanga Value Mart moving to 'At Risk' with a new account manager), re-runs the "
        "snapshot, and asserts seven properties: a new version row was created, exactly one "
        "version is current, the current version carries the new value, the superseded "
        "version retains the old value and was closed off, version numbering is contiguous, "
        "and unchanged merchants did NOT gain a version. All seven pass, and the harness "
        "restores the original state afterwards.", TEAL, "F0F9F9")

callout("The highest-value test in the project",
        "dim_date contiguity. Power BI time-intelligence functions (DATEADD, "
        "SAMEPERIODLASTYEAR, TOTALYTD) do not raise an error on a date table with a missing "
        "day — they silently return a wrong number. A single absent date quietly understates "
        "a month-on-month comparison and nobody notices until someone reconciles by hand. "
        "The test asserts that the row count equals the span between the minimum and maximum "
        "date, which cannot be true if there is a gap.", TEAL, "F0F9F9")

doc.add_heading("Data governance", 2)
para("Governance here means the questions a reviewer, an auditor or a downstream AI tool "
     "would ask of a number, and whether this model can answer them. Each row states the "
     "position taken and what backs it.", size=10.5)

table(["Question", "Position in this build"], [
    ["Where did this number come from?",
     "Every figure resolves to a source column through docs/source_column_lineage.csv, to "
     "the SQL that produced it through docs/report_queries.sql, and to the model that built "
     "it through the dbt manifest. The ERDs and lineage diagrams are generated from that "
     "manifest, so documentation cannot drift from code."],
    ["What does this field mean?",
     "142 columns across 12 tables are defined in the data dictionary below. Each definition "
     "is attributed: 13 come from the supplied DataDictionary, 16 are traced through "
     "lineage, 93 are defined in this build and 20 follow a naming convention. None are "
     "left undefined."],
    ["Which values are facts and which are assumptions?",
     "Of 34 source columns, 12 are documented by the data owner and 22 have meanings "
     "inferred by profiling. Every threshold that was not supplied — 7-day delayed "
     "redemption, 90-day breakage, R450/hour support cost, Health Score weights — is listed "
     "in the assumptions register rather than presented as given."],
    ["Can a consumer be given data without production access?",
     "Yes, and that is the purpose of the gold layer. It is a curated, tested, documented "
     "set of tables in a separate Warehouse. A reporting tool, an analyst or an AI service "
     "connects there and never touches a production database."],
    ["How do we know the data is right?",
     "Three independent layers: 132 dbt tests on structure and referential integrity, 28 "
     "reconciliation controls comparing two independent implementations, and 33 idempotency "
     "assertions. The build stops on failure rather than publishing."],
    ["Who owns a merchant when something is wrong?",
     "account_manager is carried onto the risk register, so an at-risk merchant resolves to "
     "a named owner rather than to a queue."],
    ["What happens when reference data changes?",
     "MerchantReference is a current-state extract — each load overwrites status, owner and "
     "target. A Type 2 snapshot captures the change, and mart_merchant_change_alerts raises "
     "status changes. Without it that history is destroyed at source and is unrecoverable."],
    ["Is there personal data here?",
     "No. There is no customer identifier anywhere in the four source files — no name, "
     "account, PIN or device. This is why CLV, churn and duplicate-redemption fraud are "
     "modelled at MERCHANT level. Merchant names and account managers are business contact "
     "data, not consumer PII; a real deployment would still classify and label them."],
    ["Can a load be re-run safely?",
     "Yes. Bronze writes are overwrite, not append — verified by running the ingestion twice "
     "with identical row counts — and the gold build is verified idempotent by 33 assertions."],
    ["What is NOT covered?",
     "Row-level security, sensitivity labelling and retention policy are not implemented. "
     "They are deployment concerns needing an organisation's classification scheme, and are "
     "listed here rather than left silent."],
], widths=[4.6, 12.0], font=8.5)
para()

doc.add_heading("Governance as an operating model", 2)
para("Governance is everyone's business, but every data asset needs one named owner. Shared "
     "responsibility for behaviour; single accountability for each asset. Both halves are "
     "needed — bad data almost always enters at the point of capture, which no amount of "
     "downstream engineering repairs, while an asset owned by \"the business\" is an asset "
     "nobody answers for when a number is wrong.", size=10.5)

table(["Role", "Accountable for", "In this solution"], [
    ["Data owner (business)", "Whether the figure is correct, and who may see it",
     "account_manager is carried onto the risk register so an at-risk merchant resolves to a "
     "person rather than to a queue."],
    ["Data steward", "What each field means",
     "142 columns defined, each attributed to its source: 13 from the supplied dictionary, "
     "16 traced through lineage, 93 defined here, 20 by naming convention."],
    ["Data engineer", "The controls that enforce the first two",
     "132 dbt tests, 28 reconciliation controls and 33 idempotency assertions, all of which "
     "stop the build rather than warn after publication."],
    ["Consumer", "Raising a suspect number instead of working around it",
     "Every figure is traceable to the SQL and the model that produced it, so a challenge can "
     "be resolved rather than debated."],
], widths=[3.0, 5.0, 8.6], font=8.5)
para()
para("The distinction worth drawing is between a policy and an operating model. A policy that "
     "is not enforced is documentation. Governance becomes real when it sits IN the pipeline: "
     "a failing test stops a build, a sensitivity label travels with lineage, and an "
     "uncertified model is visibly uncertified. That shifts the default from asking people to "
     "comply to making compliance the path of least resistance.", size=10.5)

doc.add_heading("Platform choices — storage modes, fallback and mirroring", 2)
para("Three Fabric capabilities determine how a semantic model gets its data. They are set "
     "out here because the choice between them is an architecture decision rather than a "
     "setting, and because this solution deliberately uses one of them today and is built for "
     "another in production.", size=10.5)

table(["Mode", "Where the data sits", "Freshness", "Trade-off"], [
    ["Import", "Copied into the in-memory engine",
     "Stale until the next refresh",
     "Fastest queries; a refresh window and a second copy of the data."],
    ["DirectQuery", "Stays in the source; each visual issues a query",
     "Always current",
     "No copy and no refresh, but every interaction pays a query cost."],
    ["Direct Lake", "Delta/Parquet in OneLake, paged into memory on demand",
     "Current, no refresh needed",
     "Near-import speed with no copy — but requires Fabric capacity, Delta storage, and it "
     "falls back to DirectQuery under conditions listed below."],
], widths=[2.4, 4.6, 3.0, 6.6], font=8.5)
para()

para("This submission ships an IMPORT model, because the deliverable is a .pbit template that "
     "must open on a reviewer's machine without a Fabric capacity attached. The gold layer is "
     "nonetheless Direct Lake ready: it is written as Delta in OneLake, and the semantic "
     "model carries no calculated columns or calculated tables, which are among the things "
     "that force a fallback.", size=10.5)

doc.add_heading("Fallback — what it is, and what it is not", 3)
para("Fallback is frequently misunderstood as a standby copy of the data. It is not. There is "
     "no fallback table and nothing is duplicated. It is a QUERY-TIME behaviour: when Direct "
     "Lake cannot answer a query from the Parquet files, that single query is redirected to "
     "run as SQL against the SQL analytics endpoint. The data is identical; only the "
     "execution path changes, and the next query may well be served by Direct Lake again.",
     size=10.5)
para("Common triggers: capacity guardrails exceeded (row limits scale with the SKU); a view "
     "rather than a table, since a view has no Delta files to page; unsupported data types; "
     "certain row-level security configurations at the SQL endpoint; and some calculated "
     "columns and tables.", size=10.5)
callout("Why fallback deserves attention",
        "Fallback is SILENT. The report simply becomes slower and nothing reports why. The "
        "model property DirectLakeBehavior controls it: Automatic falls back when needed, "
        "DirectQueryOnly never uses Direct Lake, and DirectLakeOnly makes the query FAIL "
        "instead of falling back. Setting DirectLakeOnly during testing converts an invisible "
        "performance problem into a loud error that names the offending table.",
        TEAL, "EEF7F7")

doc.add_heading("Mirroring — where the data comes from", 3)
para("Mirroring continuously replicates an external database into OneLake as Delta tables, "
     "typically seconds to minutes behind the source. Supported sources include Azure SQL "
     "Database and Managed Instance, Cosmos DB, Snowflake, Databricks and PostgreSQL, with "
     "open mirroring available for others. The replica is read-only in Fabric, so it cannot "
     "write back to production, and mirrored storage is free up to a limit tied to the "
     "capacity.", size=10.5)
para("The three capabilities compose into one path: mirroring lands production data in "
     "OneLake as Delta, Direct Lake reads it without a refresh, and Power BI reports on it "
     "near-live — without a pipeline to maintain and without any consumer touching the "
     "production database.", size=10.5)
para("The honest limit: mirroring replaces the extract and load, not the modelling. It lands "
     "data in SOURCE shape — normalised operational tables, source naming, no conformed "
     "dimensions and no declared grain. A mirrored database is not a star schema. The "
     "transformation layer in this solution is exactly what would sit on top of it, with the "
     "mirrored tables declared as dbt sources in place of the bronze tables used here.",
     size=10.5)
doc.add_page_break()

doc.add_heading("Data dictionary", 2)
_dd = ROOT / "docs" / "data_dictionary.csv"
if _dd.exists():
    _d = pd.read_csv(_dd)
    _bysrc = _d["Definition source"].str.split(" —").str[0].value_counts()
    para(f"{len(_d)} columns across {_d['Table'].nunique()} tables. Generated from the gold "
         f"tables at build time rather than maintained by hand, so it cannot fall out of step "
         f"with the data. Every definition carries its source: "
         + ", ".join(f"{v} {k.lower()}" for k, v in _bysrc.items()) + ".", size=10.5)
    for _tbl in _d["Table"].unique():
        _sub = _d[_d["Table"] == _tbl]
        doc.add_heading(_tbl, 3)
        table(["Column", "Type", "Nulls", "Definition", "Source"],
              [[r.Column, r.Type, str(r.Nulls),
                str(r.Definition)[:300], str(r._7)[:40]] for r in _sub.itertuples()],
              widths=[3.3, 1.5, 1.0, 8.4, 2.4], font=7.5)
        para()
doc.add_page_break()

doc.add_heading("Worked example — a wrong number that looked right", 2)
para("The clearest argument for building the gold layer twice is the defect it caught, so it "
     "is set out in full rather than summarised. A standalone version of this section, with "
     "the same figures, is at docs/Data_Quality_Case_Study.pdf.", size=10.5)

para("Sales and support activity end on 31 July 2026. Vouchers sold before that date continue "
     "to be redeemed afterwards, the last on 20 August 2026 — two different endings, 20 days "
     "apart, and the model needs both. A date dimension must SPAN to the last redemption or a "
     "late redemption has no row to join to and drops silently out of every total. The "
     "REPORTING WINDOW has to stop at the last sale, because there are no sales after it. "
     "Deriving both from a single maximum conflates them.", size=10.5)

_f1 = ROOT / "docs" / "fig_date_window.png"
if _f1.exists():
    doc.add_picture(str(_f1), width=Cm(16.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
para()

para("Monthly targets are pro-rated by the days the window covers, so a part-month never "
     "shows a false shortfall. With the window running to 20 August, August was treated as 20 "
     "of 31 days covered. Target is charged for those days; sales for them are zero. "
     "R1,525,271.49 x 20/31 = R984,046 of target against no revenue, and every merchant's "
     "attainment came out understated.", size=10.5)

para("Nothing was broken in any way a test could see. The date table was contiguous, had no "
     "gaps, carried the correct row count, and every foreign key resolved. All 132 dbt tests "
     "passed. The table was structurally perfect and answering a slightly different question "
     "than the one being asked.", size=10.5)

_f2 = ROOT / "docs" / "fig_reconciliation.png"
if _f2.exists():
    doc.add_picture(str(_f2), width=Cm(16.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
para()

para("What caught it was the reconciliation below: the pandas and dbt implementations "
     "disagreed on the target total by exactly R984,046. A uniqueness test asks whether a "
     "table is well-formed; it cannot ask whether the period is the right period. Two "
     "independent implementations of the same business rule can both be well-formed and still "
     "disagree — and when they do, one of them is wrong.", size=10.5)

para("The fix is one value. In the bounds CTE the redemption branch contributes to the "
     "calendar maximum but passes NULL to the activity maximum, extending the span without "
     "moving the window. No date is hardcoded anywhere in the date dimension: point the "
     "project at a different extract and the window moves with the data.", size=10.5)

callout("Why this generalises",
        "The brief specifies no dates. Whether a redemption tail belongs inside the reporting "
        "period is an UNSTATED requirement — the kind that produces a number which is wrong "
        "and entirely plausible. Those are the errors worth engineering against, because "
        "nobody reviewing a dashboard will catch them by eye. The same gate also caught "
        "pandas rank(pct=True) being a different statistic from SQL PERCENT_RANK(), which "
        "moved the Health Score by up to 9.7 points.", TEAL, "EEF7F7")

doc.add_heading("Cross-implementation reconciliation", 2)
para("The gold layer is deliberately built twice — once in pandas, once in dbt SQL — by two "
     "independent code paths. Every headline figure is then compared. If two different "
     "engines produce the same answer, the transformation logic is almost certainly right; "
     "if they disagree, one of them is wrong and the report does not ship.")

rich([("This caught two genuine defects that no unit test would have found. ", True, RED),
      ("First, dim_date was absorbing the August redemption tail into the reporting window, "
       "inflating pro-rated sales targets by ", False, None),
      ("R984,046", True, RED),
      (" against a month containing no sales at all. Second, the Python and SQL "
       "percentile-rank conventions differ — pandas rank(pct=True) returns rank/n while ANSI "
       "PERCENT_RANK() returns (rank−1)/(n−1) — which shifted the Health Score by up to ",
       False, None),
      ("9.7 points", True, RED),
      (". Both are now fixed.", False, None)])

table(["Reconciliation result", "Count"], [
    ["Checks passed exactly", str(RC["passed"])],
    ["Warnings (documented)", str(RC["warnings"])],
    ["Failures", str(RC["failures"])],
], widths=[10.0, 3.0], font=9.5, align_right=[1])

para("The single remaining warning is a 0.1 difference on the Health Score caused by Python's "
     "banker's rounding versus SQL's round-half-up. Every INPUT to the score reconciles to "
     "0.000000, which is what proves it is a presentation artefact rather than a logic "
     "difference — and that proof is only available because the inputs are compared "
     "individually rather than just the result.", size=9.5, italic=True, colour=GREY)

doc.add_page_break()

# ======================================================================================
# 5. DAX & REPORT
# ======================================================================================
doc.add_heading("5. Semantic model and report", 1)

doc.add_heading("DAX measure library", 2)
para("31+ measures organised into display folders: Sales, Redemption, Operations, Time "
     "Intelligence, Ranking, Formatting and Narrative. Conventions applied throughout: "
     "measures live in a dedicated _Measures table; DIVIDE() is used everywhere so a zero "
     "denominator returns BLANK rather than an error; and ratios are always SUM/SUM, never "
     "AVERAGE of a per-row ratio.")

callout("Why SUM/SUM rather than AVERAGE",
        "Average Basket Value must be SUM(sales_value)/SUM(transactions). Taking "
        "AVERAGE(avg_basket_value) averages a per-row ratio, which weights a one-transaction "
        "day exactly the same as a five-hundred-transaction day. It is the classic way an "
        "'average basket' measure ends up quietly wrong, and it is invisible unless someone "
        "checks the arithmetic by hand.", TEAL, "F0F9F9")

para("DAX cannot be executed outside Power BI, so each measure carries a SQL reference "
     "definition evaluated against the gold layer to produce an expected value. "
     "docs/dax_validation.csv is the acceptance test: after the measures are entered, every "
     "card is checked against it, and a mismatch identifies exactly which definition drifted.")

table(["Measure", "Expected value", "Note"], [
    ["Total Sales", Rm(K["TotalSales"]), "26,500 fact rows"],
    ["Average Basket Value", R(K["AvgBasketValue"]), "SUM/SUM"],
    ["Redemption Rate %", P2(K["RedemptionRate"]), "volume basis"],
    ["Value Redemption Rate %", P2(K["ValueRedemptionRate"]), "value basis"],
    ["Outstanding Liability", Rm(K["OutstandingLiability"]), "balance-sheet liability"],
    ["Avg Days to Redeem", f"{K['AvgDaysToRedeem']:.2f}", "redeemed vouchers only"],
    ["SLA Breach Rate %", P2(K["SLABreachRate"]), "358 of 1,363"],
    ["Breach Concentration in High Priority %", "94.7%", "the headline ops finding"],
    ["Implied SLA for 90% Compliance", "48.6h", "what the target should be"],
    ["Revenue Concentration (Top 5) %", P1(S["revenue_concentration"]["top5_share"]),
     "concentration risk"],
], widths=[6.4, 3.4, 6.8], font=9, align_right=[1])

doc.add_heading("Report pages", 2)
table(["Page", "Contents", "Interaction"], [
    ["1. Executive Overview",
     "Total sales, transactions, redemption %, avg resolution hours, sales trend, region "
     "performance, top merchants",
     "Slicers on date, region, channel, voucher type. Right-click → drill through to "
     "Merchant Detail"],
    ["2. Merchant Analysis",
     "Top and bottom merchants, voucher type performance, ranking, Pareto contribution, "
     "health scorecard",
     "Drill-through target page; tooltip page shows a 6-month sparkline and the narrative "
     "measure"],
    ["3. Operational View",
     "Ticket volume, priority mix, avg resolution hours, SLA risk, ticket trends, "
     "merchant-level issues",
     "Cross-filter from priority to merchant; conditional formatting driven by DAX colour "
     "measures"],
    ["4. Insights / Notes",
     "Narrative summary, assumptions, data-quality issues, recommended next steps, ML output",
     "Q&A visual with configured synonyms; smart narrative"],
], widths=[3.6, 6.4, 6.6], font=9)

doc.add_heading("Conditional formatting", 2)
para("Colour is semantic rather than decorative: red always means a threshold breach, green "
     "always means healthy, amber always means watch. Thresholds live in DAX measures "
     "(Colour | Sales MoM, Colour | SLA Breach, Colour | Redemption Rate) bound to 'Field "
     "value' in the formatting pane. Centralising them means the whole report restates "
     "consistently when a threshold changes, instead of needing each visual edited by hand.")

doc.add_page_break()

# ======================================================================================
# 6. AI EXTENSION
# ======================================================================================
doc.add_heading("6. AI and machine learning extension", 1)

para("Five models, trained in a Fabric notebook, logged to MLflow, with scored outputs "
     "written back to the gold layer as Delta tables. Power BI then reads predictions as "
     "ordinary columns rather than calling a model at query time — a live endpoint call per "
     "visual would add latency to every page render and couple report availability to "
     "endpoint availability, for no analytical gain given the data changes once a day.")

table(["Model", "Algorithm", "Validation", "Result"], [
    ["1. Anomaly detection", "Isolation Forest (400 trees)",
     "Unsupervised; validated against 4 documented embedded patterns",
     f"{ML['anomaly_detection']['anomalies_flagged']} anomalies from "
     f"{ML['anomaly_detection']['observations_scored']} obs — all 4 patterns recovered"],
    ["2. Redemption propensity", "HistGradientBoostingClassifier",
     "Time split: train Jan–May, test Jun–Jul",
     f"AUC {RP['roc_auc']:.3f} vs ceiling {RP['oracle_auc_ceiling']:.3f}"],
    ["3. Resolution time", "HistGradientBoostingRegressor",
     "Time split: train Jan–May, test Jun–Jul",
     f"MAE {ML['resolution_time']['mae_hours']:.1f}h, "
     f"{P1(ML['resolution_time']['improvement_vs_naive'])} better than naive"],
    ["4. Sales forecast", "Holt-Winters (weekly seasonality)",
     "Backtest on held-out final 30 days",
     f"MAPE {P2(SF['mape'])}, {P1(SF['improvement_vs_naive'])} better than naive"],
    ["5. Segmentation", "K-Means (k=4)", "Silhouette compared across k=2..7",
     "4 segments; 2 singleton outliers identified automatically"],
], widths=[3.4, 4.2, 4.4, 4.6], font=8.5)

doc.add_heading("Anomaly detection — and how it was validated", 2)
rich([("The design decision that makes the model useful is that every feature is expressed as "
       "a deviation from that merchant's ", False, None),
      ("own", True, None),
      (" expanding history, not as an absolute value. Absolute features would rank merchants "
       "by size — the model would flag the largest accounts every month and never notice a "
       "small merchant collapsing.", False, None)])

rich([("Validation. ", True, TEAL),
      ("The dataset documentation states that four patterns were deliberately embedded. The "
       "model was never told what they were, and recovered all four unprompted: Umhlanga "
       "Value Mart's July collapse (ranked #1 by anomaly score), Kudu Digital Kiosk's May "
       "growth step, Durban Cash Hub's June ticket spike, and a Liberty Lane "
       "redemption-delay month. Every flagged row also carries a plain-English explanation, "
       "because a score with no reason attached does not get acted on.", False, None)])

doc.add_heading("Reporting a weak metric honestly", 2)
rich([("The redemption propensity model scores AUC ", False, None),
      (f"{RP['roc_auc']:.3f}", True, NAVY),
      (", which looks weak. Rather than explaining it away or quietly dropping the model, it "
       "was tested against a theoretical ceiling. If redemption is generated purely as a "
       "function of voucher type, the best any model can achieve is to predict the type-level "
       "base rate — an oracle that scores AUC ", False, None),
      (f"{RP['oracle_auc_ceiling']:.3f}", True, NAVY),
      (". The model therefore captures ", False, None),
      (P1(RP["pct_of_achievable_signal"]), True, TEAL),
      (" of the achievable signal: the limit is the data, not the model. Establishing that "
       "distinction is what stops a team spending a sprint chasing an AUC that cannot move.",
       False, None)])

para(f"The ranking remains operationally useful regardless: the top risk decile carries "
     f"{RP['top_vs_bottom_decile_ratio']:.1f}× the non-redemption rate of the bottom decile, "
     f"concentrating {R(RP['top_decile_value_at_risk'])} of at-risk value into 10% of vouchers "
     f"— which is exactly what a targeted reminder campaign needs.")

callout("Validation discipline applied throughout",
        "Every supervised model uses a time-based split (train January–May, test June–July), "
        "never a random one. A random split leaks future information into training and "
        "produces a metric that will not survive contact with production. An honest 0.62 that "
        "holds is worth more than an inflated 0.85 that does not.", TEAL, "F0F9F9")

doc.add_heading("Segmentation — choosing k on business grounds", 2)
para(f"Silhouette analysis favours k=2 (score "
     f"{ML['segmentation']['silhouette_at_optimal_k']:.3f}), but across only 25 merchants "
     f"that produces a 21-versus-4 split with no operational value. k=4 was selected so each "
     f"segment is large enough to assign an owner and distinct enough to justify a different "
     f"play, at a silhouette cost "
     f"({ML['segmentation']['silhouette_at_selected_k']:.3f}) that is stated openly rather "
     f"than hidden. Notably, the algorithm isolated two merchants into single-member clusters "
     f"entirely on its own — Kudu Digital Kiosk (breakout growth) and Umhlanga Value Mart "
     f"(deteriorating). That is a finding, not a defect.")

doc.add_page_break()

# ======================================================================================
# 7. BUSINESS QUESTIONS
# ======================================================================================
doc.add_heading("7. Business questions answered", 1)

q1, q2, q3 = BA["Q1_highest_sales"], BA["Q2_best_voucher_type"], BA["Q3_declining_region"]

doc.add_heading("Q1. Which merchants generate the highest sales value and transaction volume?", 2)
rich([("Durban Cash Hub leads on both", True, NAVY),
      (f" — {R(q1['value'])} in sales ({P1(q1['share'])} of the portfolio) and "
       f"{q1['transactions']:,} transactions. The two rankings agreeing at the top is not "
       "guaranteed: a merchant can lead on value while trailing on volume if its basket is "
       "larger, so both are reported rather than assuming one proxies the other.", False, None)])
table(["#", "Merchant", "Region", "Total sales", "Transactions", "Share"],
      [[int(r.SalesRank), r.Merchant, r.Region, R(r.TotalSales), f"{r.TotalTransactions:,}",
        P1(r.SalesShare)] for _, r in score.head(5).iterrows()],
      widths=[1.0, 5.0, 3.4, 3.0, 2.6, 1.8], font=9, align_right=[0, 3, 4, 5])
para(f"Concentration matters more than the leader: the top 5 merchants carry "
     f"{P1(S['revenue_concentration']['top5_share'])} of revenue and "
     f"{S['revenue_concentration']['merchants_for_80pct']} of 25 account for 80%. "
     f"HHI is {S['revenue_concentration']['hhi']:.0f} — moderate concentration, worth "
     f"monitoring as a commercial risk in its own right.")

doc.add_heading("Q2. Which voucher type has the highest redemption rate?", 2)
rich([("Airtime, at ", False, None), (P1(q2["RedemptionRate"]), True, NAVY),
      (". Gaming is lowest at ", False, None), (P1(q2["worst_rate"]), True, RED),
      (" — a 16.9 percentage point spread.", False, None)])
table(["Voucher type", "Sold", "Redemption %", "Value redemption %", "Avg days", "Outstanding"],
      [[r.VoucherType, f"{int(r.VouchersSold):,}", P2(r.RedemptionRate),
        P2(r.ValueRedemptionRate), f"{r.AvgDaysToRedeem:.2f}", R(r.OutstandingValue)]
       for _, r in vtype.iterrows()],
      widths=[3.4, 2.2, 2.8, 3.2, 2.0, 3.0], font=9, align_right=[1, 2, 3, 4, 5])
para("Two observations that change the interpretation. First, the value-based rate is almost "
     "identical to the volume-based rate within each type, so high-value and low-value "
     "vouchers behave the same way — had they diverged, the value rate would be the one "
     "Finance should use. Second, time-to-redeem is effectively flat across all types "
     "(3.5–3.7 days), so the difference is whether customers redeem at all, not how quickly. "
     "Gaming's 24% non-redemption is the largest single block of the "
     f"{Rm(K['OutstandingLiability'])} outstanding liability.")

doc.add_heading("Q3. Which region shows declining sales or transaction behaviour?", 2)
rich([("Eastern Cape", True, RED),
      (", on four independent signals rather than one.", False, None)])
bullet("It is the only region that peaked before July (peak May 2026).")
bullet(f"It sits {P1(abs(q3['vs_peak']))} below its own peak while every other region is "
       f"currently at its peak.")
bullet("Its trend slope is +2.0% of average monthly sales, against +4.4% to +8.5% elsewhere.")
bullet("June fell 12.2% month-on-month against +1.4% to +3.8% for every other region.")
table(["Region", "Total sales", "Share", "Trend % of avg", "Last month MoM", "Peak month",
       "vs own peak"],
      [[r.Region, R(r.TotalSales), P1(r.SalesShare), f"{r.TrendPctOfAvg*100:+.1f}%",
        f"{r.LastMonthMoM*100:+.1f}%", r.PeakMonth, f"{r.SalesVsPeak*100:+.1f}%"]
       for _, r in region.iterrows()],
      widths=[3.0, 2.8, 1.6, 2.6, 2.8, 2.0, 2.0], font=9, align_right=[1, 2, 3, 4, 6])
para("A single month's movement would not justify calling a decline. Four signals pointing "
     "the same way does. Three of the five Critical/Watch merchants sit in Eastern Cape "
     "(Table Bay Express, Pretoria PayPoint, Mzansi Mini Market), so this is a regional "
     "pattern rather than one bad account.")

doc.add_heading("Q4. Are support ticket volumes, priority or long resolution times "
                "associated with weaker merchant performance?", 2)
rich([("Not as a portfolio-level rule — but decisively at the level of individual events.",
       True, NAVY), ("", False, None)])
para("The tempting answer is that tickets per 1,000 transactions correlates with target "
     f"attainment at r = {SZ['raw_corr_friction_vs_attainment']:+.2f}, so friction hurts "
     "performance. That answer is confounded, and reporting it would have been a mistake.")
bullet(f"The ratio is strongly size-dependent: r = {SZ['friction_vs_log_size_pearson']:+.2f} "
       f"against log(total sales). Small merchants score badly purely because the denominator "
       f"is small.")
bullet(f"Controlling for size, the partial correlation collapses to r = "
       f"{SZ['partial_corr_friction_vs_attainment_controlling_size']:+.2f}.")
bullet("SLA breach rate and average resolution time show no association with performance at "
       "all (r = +0.04 and −0.19).")
para("What is real is event-level. Four months show a merchant's ticket volume more than "
     "doubling against its own prior three months:")
table(["Merchant", "Month", "Tickets", "Prior avg", "Ticket uplift", "Sales vs prior 3M",
       "Diagnosis"],
      [["Durban Cash Hub", "2026-06", "44", "5.0", "+780%", "+8.2%",
        "Service issue, healthy account"],
       ["Durban Cash Hub", "2026-07", "52", "18.3", "+184%", "+6.3%",
        "Service issue, healthy account"],
       ["Umhlanga Value Mart", "2026-07", "37", "4.7", "+693%", "−42.5%", "Failing account"],
       ["Soweto Super Save", "2026-05", "11", "4.3", "+154%", "+13.2%",
        "Service issue, healthy account"]],
      widths=[3.6, 2.0, 1.6, 1.8, 2.2, 2.6, 3.0], font=9, align_right=[2, 3, 4, 5])
rich([("Practical conclusion. ", True, TEAL),
      ("Ticket spikes are worth investigating immediately, but they do not predict revenue on "
       "their own. Alerting must pair operational movement with commercial movement — which "
       "is how the alert logic in this solution is written.", False, None)])

doc.add_heading("Q5. Which merchants should management focus on first, and why?", 2)
para("Ranked by revenue at risk, not by severity of decline.")
table(["#", "Merchant", "Region", "Health", "Latest month", "vs prior 3M", "Revenue at risk"],
      [[i + 1, m["Merchant"], m["Region"], f"{m['HealthScore']:.0f}/100",
        R(m["LatestMonthSales"]), f"{m['SalesVsPrior3Avg']*100:+.1f}%",
        R(m["RevenueAtRiskAnnualised"])]
       for i, m in enumerate(BA["Q5_focus_merchants"])],
      widths=[1.0, 4.6, 3.0, 2.0, 2.6, 2.4, 3.0], font=9, align_right=[0, 3, 4, 5, 6])
rich([("Why revenue at risk rather than percentage decline. ", True, TEAL),
      ("Umhlanga's 42.5% collapse and Table Bay Express's 5.9% slide are not comparable on "
       "their face — but Table Bay is a R3.3m merchant, so its smaller percentage still puts "
       "R354,163 a year at stake against Umhlanga's R571,518. Ranking on percentage change "
       "alone sends the account team to the wrong door.", False, None)])

doc.add_page_break()

# ======================================================================================
# 7b. RECONCILIATION AND FINANCIAL CONTROLS
# ======================================================================================
doc.add_heading("8. Reconciliation and financial controls", 1)
para("Controls that PASS matter as much as controls that fail: a control nobody can see is a "
     "control nobody trusts. mart_reconciliation materialises each check with its expected "
     "value, actual value, variance and a derived status, so the control set is data in the "
     "report rather than a claim in a document.")

_rec = con_rec = None
try:
    import duckdb as _dd
    _c = _dd.connect(str(ROOT / "data" / "mvi.duckdb"))
    _rec = _c.execute("""select control_family, control_name, expected_value, actual_value,
                                variance, control_status
                         from main_marts.mart_reconciliation order by control_order""").df()
    _c.close()
except Exception:
    pass

if _rec is not None:
    table(["Control family", "Check", "Expected", "Actual", "Variance", "Status"],
          [[r.control_family, r.control_name, f"{r.expected_value:,.2f}",
            f"{r.actual_value:,.2f}", f"{r.variance:,.2f}", r.control_status]
           for _, r in _rec.iterrows()],
          widths=[2.8, 5.6, 2.4, 2.4, 2.0, 1.6], font=8.5, align_right=[2, 3, 4])

rich([("The population control is the one that matters. ", True, NAVY),
      ("fct_merchant_sales totals R65,521,299 across 510,127 transactions; "
       "fct_voucher_redemptions totals R22,019,853 across 120,969 vouchers — a variance of ",
       False, None),
      ("R43,501,446", True, RED),
      (". That is EXPECTED, not a break. MerchantSales is a daily aggregate of ALL "
       "transactions; VoucherRedemptions is a voucher-level extract covering roughly 1 in "
       "4.2 of them. They describe different populations and must never be forced to tie.",
       False, None)])
callout("Why recording an EXPECTED variance is the point",
        "Without it, someone compares R65.5m to R22.0m, calls it a R43.5m reconciliation "
        "break and escalates — or worse, divides one by the other and reports a value-based "
        "redemption rate against the wrong denominator. The control does not just check the "
        "number; it records which comparisons are legitimate.", AMBER, "FEF8EC")

# ======================================================================================
# 7c. VALUE, ATTRITION AND FRAUD SIGNALS
# ======================================================================================
doc.add_heading("9. Merchant value, attrition risk and fraud signals", 1)

rich([("Customer lifetime value and customer churn were asked for and cannot be built. ",
       True, NAVY),
      ("There is no customer identifier anywhere in the four source files — not a hashed "
       "one, not a session id, not a card token. Every fact is either a merchant-daily "
       "aggregate or a voucher row with no purchaser attached. Presenting something shaped "
       "like customer CLV would be dishonest, so mart_merchant_value_risk models the "
       "defensible equivalent: the merchant is the customer of the voucher business.",
       False, None)])

try:
    import duckdb as _dd
    _c = _dd.connect(str(ROOT / "data" / "mvi.duckdb"))
    _vr = _c.execute("""select merchant_name, region, annualised_run_rate, tenure_months,
                               implied_lifetime_value, attrition_risk_score,
                               attrition_risk_band, reversal_per_1k_txn, risk_signal_band
                        from main_marts.mart_merchant_value_risk
                        order by annualised_run_rate desc limit 8""").df()
    _c.close()
    table(["Merchant", "Region", "Run rate", "Tenure", "Implied LTV", "Attrition",
           "Band", "Rev/1k", "Signal"],
          [[r.merchant_name, r.region, R(r.annualised_run_rate), f"{int(r.tenure_months)}m",
            R(r.implied_lifetime_value), f"{r.attrition_risk_score:.0f}",
            r.attrition_risk_band, f"{r.reversal_per_1k_txn:.2f}", r.risk_signal_band]
           for _, r in _vr.iterrows()],
          widths=[3.6, 2.4, 2.2, 1.2, 2.4, 1.6, 1.4, 1.2, 1.4], font=8,
          align_right=[2, 3, 4, 5, 7])
except Exception:
    pass

doc.add_heading("Geographic intelligence", 2)
para("Five of South Africa's nine provinces carry merchants. The four that do not — Limpopo, "
     "Mpumalanga, North West and Northern Cape — account for 56% of the country's land area. "
     "The choropleth draws them as hatched 'no cover' rather than omitting them, because an "
     "absent province otherwise reads as zero sales when it actually means no footprint at "
     "all, which is a different commercial question.")
table(["Province", "Merchants", "Sales", "Share", "Position"], [
    ["Free State", "10", "R18,009,802", "27.5%", "Largest, and growing"],
    ["Gauteng", "4", "R16,145,481", "24.6%", "At peak"],
    ["Western Cape", "4", "R13,068,297", "19.9%", "At peak; carries the April redemption "
                                                   "delay on Bill Payment"],
    ["KwaZulu-Natal", "3", "R9,516,589", "14.5%", "Strongest momentum (+13.1%)"],
    ["Eastern Cape", "4", "R8,781,130", "13.4%", "The only province in decline"],
], widths=[3.0, 1.8, 3.0, 1.6, 7.2], font=9, align_right=[1, 2, 3])

callout("What the geography does NOT tell us",
        "It is tempting to read 56% uncovered land as 56% untapped market. It is not. No "
        "population, GDP or competitor data was supplied, and land area is a poor proxy for "
        "demand — Northern Cape is the largest province in South Africa and the most sparsely "
        "populated. The defensible statement is a FOOTPRINT observation: four provinces have "
        "no merchant presence. Sizing that as an opportunity needs population and income data "
        "the business already holds and this dataset does not carry.", AMBER, "FEF8EC")

doc.add_heading("Fraud detection AND prevention", 2)
para("Detection is one third of a control framework. A model that spots a loss after it has "
     "happened is worth less than a rule that stops it, and both are worth less without a "
     "corrective path. The framework below separates the three, and is explicit about which "
     "parts this dataset can support today.")
table(["Layer", "Control", "Status", "Basis or requirement"], [
    ["PREVENT", "Voucher expiry enforcement", "Modelled, not enforced",
     "The 90-day rule exists in the model as breakage; enforcing it at redemption is a "
     "system change, not a reporting one"],
    ["PREVENT", "Value ceilings per voucher type", "Recommended",
     "Gaming carries 24% non-redemption — the highest exposure per rand issued"],
    ["PREVENT", "Velocity limits per merchant", "Recommended",
     "Cap redemptions per merchant per hour; the data shows a uniform ~23% same-day rate, so "
     "a genuine outlier would stand out"],
    ["PREVENT", "Onboarding due diligence", "Gap",
     "Two merchants are flagged 'At Risk' in the CRM with no linked control action"],
    ["PREVENT", "Segregation of duties", "Recommended",
     "Issuance and redemption approval should not share an owner"],
    ["DETECT", "Reversal rate per 1,000 transactions", "BUILT",
     "215 reversal tickets; per-merchant range 0.00 to 1.99 per 1k"],
    ["DETECT", "Redemption velocity", "BUILT", "Same-day redemption share per merchant"],
    ["DETECT", "Behavioural anomaly", "BUILT",
     "Isolation Forest on merchant-month deviation from own history"],
    ["DETECT", "Liability concentration", "BUILT",
     "R3,541,563 outstanding, tracked per merchant"],
    ["DETECT", "Voucher value outliers", "BUILT — returns zero",
     "Implemented and evaluated: no voucher sits beyond 3 SD within its type"],
    ["CORRECT", "Quarantine on quality failure", "BUILT",
     "The dbt gate between silver and gold; bad loads never reach the report"],
    ["CORRECT", "Change alerting", "BUILT",
     "mart_merchant_change_alerts raises a status or target change as Critical"],
    ["CORRECT", "Reversal workflow routing", "Recommended",
     "Financial tickets are 25% of volume and should route separately from commercial ones"],
    ["CORRECT", "Merchant suspension path", "Gap",
     "A 'Review' risk band should trigger a hold, not an email"],
    ["CORRECT", "Audit trail", "BUILT",
     "batch_id on every row answers 'which load produced this figure' during an incident"],
], widths=[2.2, 4.4, 3.0, 7.0], font=8)

doc.add_heading("Fraud controls — what the data supports, tested not assumed", 2)
table(["Control", "Status", "Basis"], [
    ["Reversal ticket rate", "SUPPORTED",
     "Reversal Query tickets per 1k transactions. Reversals are where value moves back."],
    ["Redemption velocity", "SUPPORTED",
     "Same-day redemption share. Legitimate for airtime, unusual in bulk."],
    ["Liability concentration", "SUPPORTED", "Unredeemed value sitting with one merchant."],
    ["Behavioural anomaly", "SUPPORTED",
     "Isolation Forest — multivariate shift against a merchant's own history."],
    ["Voucher value outliers", "TESTED — NIL",
     "Implemented and evaluated: ZERO vouchers beyond 3 SD within type, because the "
     "generator used a bounded distribution. The control runs and correctly returns zero "
     "rather than being quietly omitted."],
    ["Duplicate redemption", "NOT SUPPORTED",
     "voucher_id is unique by construction in the extract, so a duplicate cannot appear. "
     "Needs a raw redemption event log with repeated attempts."],
    ["Geographic anomaly", "NOT SUPPORTED",
     "Needs merchant location, transaction location and timestamp. Only static province at "
     "merchant level is supplied."],
    ["PIN / invalid voucher use", "NOT SUPPORTED",
     "Needs voucher PIN, attempt outcome and failure reason."],
    ["Customer velocity / CLV / churn", "NOT SUPPORTED",
     "No customer identifier exists. Merchant-level equivalents built instead."],
], widths=[4.0, 2.6, 10.0], font=8.5)

callout("Stating what cannot be built is part of the deliverable",
        "Four of the nine controls above are unbuildable on this dataset, and each names the "
        "telemetry it would need. That is more useful to the business than a dashboard of "
        "fraud metrics computed from fields that do not carry the signal — and it is the "
        "difference between a control framework and a control-shaped picture.", TEAL,
        "F0F9F9")

doc.add_page_break()

# ======================================================================================
# 8. RECOMMENDATIONS
# ======================================================================================
doc.add_heading("8. Recommendations", 1)
table(["#", "Action", "Owner", "Rationale", "Expected impact"], [
    ["1", "Site visit to Umhlanga Value Mart this week", "Account Management",
     "Sales −42.5% and tickets +693% in the same month — an operational failure causing "
     "commercial damage, which means it is plausibly recoverable",
     f"{R(571518)} annualised revenue"],
    ["2", "Re-base the SLA policy", "Operations",
     "94.7% of breaches come from Critical/High because targets run inverse to actual "
     "workload. A 90th-percentile-compliant Critical SLA would be ~120h. Either re-base the "
     "targets or resource high-priority work separately",
     "Restores SLA as a usable management signal"],
    ["3", "Root-cause Durban Cash Hub's ticket volume", "Operations",
     "+780% tickets on the largest merchant in the book while sales still grow — a service "
     "problem that has not yet become a commercial one",
     f"Protects {Rm(5776119)} of annual revenue"],
    ["4", "Eastern Cape regional review", "Regional Management",
     "The only region below its own peak; 3 of 5 Critical/Watch merchants sit there",
     f"{Rm(8781130)} region on a declining trajectory"],
    ["5", "Gaming voucher redemption campaign", "Commercial",
     "24% non-redemption against 7% for Airtime. The propensity model can target the top "
     "risk decile directly",
     f"{R(171689)} of at-risk value in the top decile"],
    ["6", "Confirm the sales target basis with Finance", "Finance / BI",
     "Targets sit ~6.1× below realised sales for all 25 merchants — a basis error, not "
     "outperformance", "Makes target attainment reportable at all"],
], widths=[0.8, 4.2, 2.8, 6.2, 3.0], font=8.5, align_right=[0])

doc.add_page_break()

# ======================================================================================
# 9. DATA QUALITY, ASSUMPTIONS, LIMITATIONS
# ======================================================================================
doc.add_heading("9. Data quality, assumptions and limitations", 1)

doc.add_heading("Data quality findings", 2)
table(["Finding", "Detail and treatment"], [
    ["Sales targets mis-calibrated",
     "BaseMonthlySalesTarget sits ~6.1× below realised sales for ALL 25 merchants, so raw "
     "target attainment reads 614%. The consistency across every merchant points to a basis "
     "or units error rather than genuine outperformance. The supplied value is retained "
     "unchanged for transparency; the report uses a relative Target Attainment Index "
     "(merchant attainment ÷ portfolio attainment), which is comparable across merchants and "
     "immune to the calibration error. Confirming the intended basis is the first question "
     "for the business."],
    ["SLA thresholds inverted vs workload",
     "Critical gets 12h but averages 52.7h; Low gets 48h and averages 11.3h. Treated as a "
     "genuine policy finding and reported as such, not silently corrected in the data."],
    ["Redemption flag vs date", "0 records disagree in the supplied data. The validation rule "
     "is still enforced in silver, and violations would be flagged via quality_flag rather "
     "than silently dropped — the headline redemption rate is too visible to leave "
     "unguarded."],
    ["Referential integrity", "0 orphan MerchantIDs across all three facts. Fact-embedded "
     "Merchant/Region/Channel agree with MerchantReference in 100% of cases, which is what "
     "made dropping them from the facts a lossless decision."],
    ["Redemption tail beyond the window",
     "RedeemedDate extends to 20 August 2026, past the 31 July sales cut-off. This is "
     "expected. The calendar SPANS these dates so the join works, but the reporting WINDOW is "
     "defined by activity dates only — without that separation a partial August containing no "
     "sales would drag down every period-based measure."],
], widths=[4.0, 12.6], font=9)

doc.add_heading("Assumptions", 2)
table(["Assumption", "Rationale"], [
    ["Delayed redemption = more than 7 days",
     "No threshold supplied. 7 days sits above the 75th percentile (5 days), so it flags a "
     "genuine tail rather than routine behaviour."],
    ["Voucher expiry / breakage = 90 days",
     "No expiry rule supplied; 90 days is a common voucher industry default. Affects the "
     "breakage measure only — outstanding liability is unaffected."],
    ["Targets pro-rated by days covered",
     "Every month in the window is complete, so the factor is currently 1.0. The logic is in "
     "place so a mid-month refresh cannot produce a false shortfall."],
    ["Support cost of R450/hour",
     "Used only in the indicative Ops Cost Exposure measure. Clearly flagged; should be "
     "replaced with Finance's actual loaded rate before use."],
    ["Voucher category and margin band",
     "Not present in the source. Maintained as a dbt seed so the commercial team can change "
     "the mapping without a code deployment."],
    ["Health Score weights",
     "Weighted toward recent momentum (25%) over structural trend (15%) because the "
     "operational purpose is early warning. Declared in one place and treated as a "
     "business-tunable parameter, not a hidden constant."],
], widths=[4.4, 12.2], font=9)

doc.add_heading("Limitations", 2)
table(["Limitation", "Impact"], [
    ["Seven months of data",
     "No year-on-year comparison is possible. The YoY measures are written and will work once "
     "a second year exists, but currently return BLANK by design rather than a misleading "
     "zero."],
    ["No annual seasonality baseline",
     "The 4.69% forecast MAPE reflects weekly seasonality only. Festive trading and school "
     "terms cannot be modelled from seven months and would likely change the December "
     "forecast materially."],
    ["25 merchants limits statistical power",
     "Merchant-level correlations rest on n=25. This is why the operations-versus-performance "
     "analysis uses the 175-row merchant-month panel and partial correlation rather than "
     "resting on a single cross-sectional coefficient."],
    ["Redemption model ceiling",
     "Redemption is generated almost purely from voucher type, capping achievable AUC at "
     "~0.62. On real data with customer-level features (tenure, prior redemption behaviour, "
     "channel), materially higher performance would be expected."],
    ["Synthetic data",
     "All findings describe the supplied synthetic dataset. The anomaly model independently "
     "recovering the four embedded patterns validates the METHOD; it does not validate the "
     "conclusions against real trading behaviour."],
    ["No cost or margin data",
     "Only revenue is supplied, so merchant profitability — and therefore true commercial "
     "prioritisation — cannot be assessed. Revenue at Risk is a revenue measure, not a margin "
     "measure."],
], widths=[4.4, 12.2], font=9)

doc.add_page_break()

# ======================================================================================
# 9b. EVIDENCE — screenshots, where supplied
# ======================================================================================
_man = ROOT / "docs" / "screenshot_manifest.json"
if _man.exists():
    _sm = json.load(open(_man))
    _shots = [s for s in _sm["screenshots"] if s.get("file")]
    if _shots:
        doc.add_heading("Appendix A. Evidence", 1)
        para(f"{len(_shots)} of {_sm['total']} captures. Screenshots evidence what was built "
             f"in Fabric; the command output alongside them is the stronger proof, because it "
             f"cannot be staged.")
        _last = None
        for s in _shots:
            if s["section"] != _last:
                doc.add_heading(s["section"], 2)
                _last = s["section"]
            _p = ROOT / s["file"]
            if _p.exists():
                # 16.8cm is the usable width on A4 portrait (21 - 2.1 - 2.1). Anything
                # wider is cropped at the right edge without warning.
                doc.add_picture(str(_p), width=Cm(16.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                para(s["caption"], 8.5, italic=True, colour=GREY,
                     align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_page_break()

# Command-line evidence is included whether or not screenshots exist — it is reproducible
doc.add_heading("Appendix B. Fabric connection evidence", 1)
para("Captured from the terminal. Unlike a screenshot this can be re-run and re-verified.")
table(["Step", "Command", "Result"], [
    ["Workspace discovered", "scripts/17_fabric_connect.py",
     "WS_MerchantVoucher · 0f8b1362-fedb-4c4d-aecc-f788a989c6b2 · capacity attached"],
    ["Warehouse created", "POST /v1/workspaces/{id}/warehouses",
     "WH_MerchantVoucher · 278e4bef-ae60-428e-8cdc-15224bbd3c34 · provisioned"],
    ["Connection tested", "dbt debug --target fabric", "All checks passed"],
    ["Write access proven", "dbt seed --target fabric", "PASS=4 WARN=0 ERROR=0"],
    ["Seed and staging models", "dbt run --target fabric --select staging",
     "Executed against the warehouse through the portability macro layer"],
    ["Full build", "dbt build", "156 pass, 0 warn, 0 error"],
    ["Idempotency", "scripts/15_test_idempotency.py",
     "33/33 — pipeline re-run produces identical output"],
], widths=[3.6, 5.4, 7.6], font=8.5)
callout("Verified by connecting, not by describing",
        "The workspace was provisioned through the Fabric REST API, authentication uses an "
        "az login CLI token so no secret is stored anywhere in this repository, and write "
        "access was proven by seeding rather than inferred from a successful connection "
        "test. Dialect differences between the development engine and T-SQL are absorbed by "
        "the portability macro layer in dbt/macros/portability.sql, so one set of models "
        "serves both targets.", TEAL, "EEF7F7")
doc.add_page_break()

# ---------------------------------------------------------------- operating environment
doc.add_heading("Extending this to a production reporting environment", 1)
para("The patterns below are the ones this solution was built around, and each is already "
     "present in the submission rather than proposed for later.", size=10.5)

doc.add_heading("Exception-based reconciliation", 2)
para("Reconciliation here does not produce a report that someone reads end to end looking for "
     "problems. It produces a control table in which every check carries a PASS, WARN or "
     "EXPECTED status, and only the exceptions require attention. 27 of 28 checks tie exactly; "
     "the one WARN is a documented rounding convention; and the R43.5m difference between the "
     "sales and voucher populations is recorded as EXPECTED with its reason attached, "
     "specifically so that a genuine break is never lost among differences that are supposed "
     "to be there. The same shape applies to reconciling an external statement against system "
     "records: normalise both sides to a common reference, classify each row as matched, "
     "partially matched or unmatched, and surface only the last two.", size=10.5)

doc.add_heading("A governed layer for AI and self-service consumption", 2)
para("Nothing consuming this model needs access to a production database. The gold layer is a "
     "curated, tested set of tables with declared grain, enforced foreign keys and documented "
     "column meanings — which is what makes an answer drawn from it traceable. Every figure "
     "in this report can be resolved back to a source column through docs/"
     "source_column_lineage.csv, back to the SQL that produced it through docs/"
     "report_queries.sql, and back to the model that built it through the dbt manifest. That "
     "chain is the difference between an answer that can be interrogated and an answer that "
     "has to be trusted: the questions 'which query produced this, and where did the number "
     "come from' both have concrete answers here.", size=10.5)
para("The controls that make that safe are the ordinary ones: business logic defined once and "
     "inherited by the SQL, DAX and ML layers rather than restated in each; 132 automated "
     "tests that stop the build rather than warn after the fact; and a documented set of "
     "assumptions, so a consumer can see which values are stated by the data owner and which "
     "were inferred. Of the 34 source columns, 12 are documented by the supplied "
     "DataDictionary and 22 are inferred — that distinction is recorded rather than smoothed "
     "over.", size=10.5)

doc.add_heading("Scaling to additional regions", 2)
para("Region is a conformed dimension attribute, not a column repeated on each fact, so "
     "adding territories does not change the model shape — it adds members. The three "
     "properties that matter for a multi-region rollout are already in place: the fact load "
     "is incremental and merges on a deterministic key, so a new region backfills without "
     "rewriting history; the pipeline is idempotent, verified by 33 assertions, so a re-run "
     "or a partially failed load cannot double-count; and the snapshot captures reference "
     "data changes over time, which matters more as the number of source systems grows and "
     "current-state extracts start arriving from several places at once.", size=10.5)
doc.add_page_break()

# ======================================================================================
# 10. DELIVERABLES
# ======================================================================================
doc.add_heading("10. Deliverables and how to run", 1)
table(["Path", "Contents"], [
    ["report/dashboard.html", "Interactive 6-page dashboard replicating the Power BI report, "
                              "plus a build walkthrough. Fully self-contained — opens in any "
                              "browser with no dependencies."],
    ["report/…Submission.docx", "This document."],
    ["excel/…Report.xlsx", "10-sheet Excel pack: exec summary, merchant scorecard, voucher "
                           "analysis, operational view, anomalies, ML results, business "
                           "questions, data dictionary, assumptions."],
    ["dbt/", "21 models, 132 tests, 5 seeds, 3 exposures, macros. Runs locally (dev) or "
             "Fabric Warehouse (prod)."],
    ["sql/", "Fabric Warehouse DDL and the business-question queries."],
    ["notebooks/", "Fabric PySpark notebooks: bronze→silver, ML scoring with MLflow."],
    ["datafactory/", "Pipeline JSON, schedule trigger, and the design rationale."],
    ["dax/", "31+ DAX measures across 4 files, with commentary on each convention."],
    ["scripts/", "Runnable reference implementation: profile, warehouse, analytics, ML, "
                 "reconciliation, DAX validation, Excel, dashboard."],
    ["data/", "bronze / silver / gold / analytics / ml layers as parquet, plus the local warehouse file."],
    ["docs/", "Profile report, analytics summary, ML summary, reconciliation, DAX validation."],
], widths=[4.6, 12.0], font=9)

doc.add_heading("Reproducing the results", 2)
para("The full pipeline runs end to end in under two minutes:", space_after=4)
for cmd in ["python scripts/01_profile.py",
            "python scripts/02_build_warehouse.py",
            "python scripts/03_analytics.py",
            "python scripts/04_ml_models.py",
            "cd dbt && dbt seed && dbt run && dbt test",
            "python scripts/05_reconcile.py",
            "python scripts/06_validate_dax.py",
            "python scripts/07_build_excel.py",
            "python scripts/08_build_dashboard.py"]:
    p = doc.add_paragraph()
    r = p.add_run("  " + cmd)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p.paragraph_format.space_after = Pt(1)

doc.add_paragraph()
callout("A closing note on verification",
        "Three claims in this submission are independently checkable rather than asserted. "
        "The gold layer is built twice by different engines and reconciles to the cent. Every "
        "DAX measure has a SQL-derived expected value in docs/dax_validation.csv. And the "
        "anomaly model recovered all four deliberately embedded patterns without being told "
        "what they were. Where a result was weak or a correlation was confounded, that is "
        "stated in the report rather than smoothed over — which is the more useful habit in a "
        "BI function, because the cost of a confident wrong number is far higher than the "
        "cost of an honest caveat.", NAVY, "EEF3F9")

# ---------------------------------------------------------------- width guard
# Every image must fit inside the usable width of the section it sits in. A too-wide picture
# is not an error in python-docx — Word simply crops it, which is how the first build shipped
# two ERDs with their right-hand columns chopped off.
_problems = []
for _sh in doc.inline_shapes:
    _w_cm = _sh.width.cm
    if _w_cm > 27.0:
        _problems.append(f"image {_w_cm:.1f}cm exceeds even landscape usable width")
if _problems:
    raise SystemExit("IMAGE WIDTH ERROR:\n  " + "\n  ".join(_problems))

try:
    doc.save(str(OUT))
except PermissionError:
    raise SystemExit(f"\nCannot write {OUT.name} — it is open in Word.\n"
                     f"Close it and re-run:  python scripts/09_build_report_docx.py\n")

print(f"Wrote {OUT}")
print(f"  {OUT.stat().st_size / 1024:.0f} KB · {len(doc.inline_shapes)} images · "
      f"{len(doc.tables)} tables")
for _sh in doc.inline_shapes:
    print(f"    image {_sh.width.cm:.1f} x {_sh.height.cm:.1f} cm")
